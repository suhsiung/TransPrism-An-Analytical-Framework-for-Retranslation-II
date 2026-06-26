#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
app_agent.py
重譯文本分析系統 AI Agent（雙 Agent 架構）
────────────────────────────────────────────
主 Agent    (LLM1): google/gemma-4-31b-it
             → 研究情境詢問 & 最終研究報告撰寫
工作程序Agent(LLM2): google/gemma-4-26b-a4b-it
             → 文本上傳引導 / 格式驗證 / 指標計算 / 統計分析
"""

# ============================================================================
# SECTION 1: Imports & Configuration
# ============================================================================
import sys, io, os, re, json, uuid, shutil, warnings, functools, collections
import threading, zipfile, traceback

# 強制 UTF-8 輸出（避免 Windows cp950 編碼錯誤）
if sys.platform.startswith("win"):
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace", line_buffering=True)
        sys.stderr.reconfigure(encoding="utf-8", errors="replace", line_buffering=True)
    except AttributeError:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
from pathlib import Path
from datetime import datetime, timezone, timedelta

# 台灣時區（UTC+8）：確保雲端（UTC 伺服器）與本地時間一致
TW_TZ = timezone(timedelta(hours=8))
def now_tw():
    """回傳台灣時間（UTC+8），不受伺服器時區影響。"""
    return datetime.now(TW_TZ)
from collections import Counter

import numpy as np
import pandas as pd
import scipy.stats as stats
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score, classification_report
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.cluster import KMeans, AgglomerativeClustering
from sklearn.decomposition import PCA
from scipy.cluster.hierarchy import dendrogram, linkage
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import urllib.request, urllib.error
import gradio as gr

warnings.filterwarnings('ignore')
print = functools.partial(print, flush=True)

# ── 載入 .env 環境變數 ───────────────────────────────────────────────────────
try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).parent / ".env")
except ImportError:
    print("[警告] 未安裝 python-dotenv，略過 .env 載入（請執行 pip install python-dotenv）")

# ── 路徑常數 ────────────────────────────────────────────────────────────────
BASE_DIR       = Path(__file__).parent
CORPUS_INFO    = BASE_DIR / "01-Corpus" / "03-Info"
SESSIONS_DIR   = BASE_DIR / "sessions"
SESSIONS_DIR.mkdir(exist_ok=True)

# ── LLM 設定 ────────────────────────────────────────────────────────────────
LLM1_API_KEY = os.environ.get("LLM1_API_KEY", "")
LLM1_MODEL   = os.environ.get("LLM1_MODEL",   "google/gemma-4-31b-it")

LLM2_API_KEY = os.environ.get("LLM2_API_KEY", "")
LLM2_MODEL   = os.environ.get("LLM2_MODEL",   "google/gemma-4-26b-a4b-it")

ALPHA = 0.05

# ── POS 集合（與 Prog02 相同）───────────────────────────────────────────────
Vset   = ['VC','VH','VE','VJ','VA','V2','VK','VL','VD','VCL','VHC','VB','VF','VI','VAC','VG']
Nset   = ['Na','Nb','Nc','Ncd','Nd','Nf','Nh','Nv']
Nhset  = ['Nh']
ADVset = ['D','Da','Dfa','Dfb','Dk']
Cset   = ['Caa','Cbb']
Caset  = ['Caa']
Cbset  = ['Cbb']
DETset = ['Neu','Nes','Nep','Neqa']
POSTset= ['Cab','Cba','Neqb','Ng']
Tset   = ['I','T']
Aset   = ['A']

# ── 研究情境定義 ─────────────────────────────────────────────────────────────
SCENARIOS = {
    "1": {
        "title": "同一原文的人工翻譯 vs 機器翻譯",
        "short": "人機對比",
        "upload_prompt": (
            "請上傳欲分析的文本，文本檔名設計用 **A_xx.txt、B_xx.txt** "
            "作為區分人工或機器翻譯。\n"
            "上傳的文本必須使用 [ckip_segmentation_sumin](https://github.com/suhsiung/ckip_segmentation_sumin) "
            "工具完成斷詞後的 `.txt` 檔案，"
            "內容格式為：`詞_詞性`、`_詞性` 或空白行。"
        ),
        "desc": "焦點在人機差異的語言特徵。",
        "scenario_module_key": "scenario_2",
    },
    "2": {
        "title": "全譯 vs 編譯（節譯、改寫）",
        "short": "全譯vs編譯",
        "upload_prompt": (
            "請上傳欲分析的文本，文本檔名設計用 **A_xx.txt、B_xx.txt** "
            "作為區分全譯或編譯。\n"
            "上傳的文本必須使用 [ckip_segmentation_sumin](https://github.com/suhsiung/ckip_segmentation_sumin) "
            "工具完成斷詞後的 `.txt` 檔案，"
            "內容格式為：`詞_詞性`、`_詞性` 或空白行。"
        ),
        "desc": "焦點在文本完整性和改寫策略。",
        "scenario_module_key": "scenario_3",
    },
    "3": {
        "title": "不同譯者的同期譯本",
        "short": "譯者風格",
        "upload_prompt": (
            "請上傳欲分析的文本，文本檔名設計用 **A_xx.txt、B_xx.txt、C_xx.txt…** "
            "作為區分不同的譯者。\n"
            "上傳的文本必須使用 [ckip_segmentation_sumin](https://github.com/suhsiung/ckip_segmentation_sumin) "
            "工具完成斷詞後的 `.txt` 檔案，"
            "內容格式為：`詞_詞性`、`_詞性` 或空白行。"
        ),
        "desc": "焦點在個人風格差異。排除時代因素，討論譯者個人的語言偏好。",
        "scenario_module_key": "scenario_4",
    },
    "4": {
        "title": "同一原文的不同時期人工譯本",
        "short": "歷時重譯",
        "upload_prompt": (
            "請上傳欲分析的文本，文本檔名設計用 **A_xx.txt、B_xx.txt、C_xx.txt…** "
            "作為區分時期。\n"
            "上傳的文本必須使用 [ckip_segmentation_sumin](https://github.com/suhsiung/ckip_segmentation_sumin) "
            "工具完成斷詞後的 `.txt` 檔案，"
            "內容格式為：`詞_詞性`、`_詞性` 或空白行。"
        ),
        "desc": "分析焦點在歷時變化、語言現代化、翻譯規範的演變。時間維度是核心變數。",
        "scenario_module_key": "scenario_1",
    },
}

# ── 指標分類（同 prog03）────────────────────────────────────────────────────
INDEX_CATEGORIES = {
    "基礎統計": {
        "AA_01": "總字數", "AA_02": "相異字數", "AA_03": "總詞數",
        "AA_04": "相異詞數", "AA_05": "段落數", "AA_06": "總句數",
    },
    "詞彙難易度": {"AC_01": "非常用詞比率", "AC_02": "難詞比率"},
    "詞彙多樣性": {"AD_02": "MSTTR", "AD_03": "RTTR"},
    "語意複雜度": {
        "AG_01": "詞彙密度1", "AG_02": "詞彙密度2",
        "AG_05": "每千詞成語數1", "AG_06": "每千詞成語數2",
    },
    "語法複雜度": {
        "AH_01": "平均句長", "AH_03": "句長標準差",
        "AH_06": "平均小句長度", "AH_08": "小句標準差",
    },
    "篇章連貫": {
        "AJ_01": "對等關聯連接詞密度",
        "AJ_02": "連接詞總密度",
        "AJ_03": "從屬連接詞密度",
        "AJ_04": "指示代詞密度",
        "AJ_05": "一般代名詞密度",
        "AJ_06": "指示照應代詞密度",
        "AJ_07": "第一人稱代詞密度",
        "AJ_08": "第二人稱代詞密度",
        "AJ_09": "第三人稱代詞密度",
        "AJ_10": "泛指與反身代詞密度",
        "AJ_11": "泛指與反身代詞密度2",
    },
    "詞性標記": {
        "AF_01": "A非謂形容詞", "AF_02": "Caa對等連接詞", "AF_03": "Cab連接詞",
        "AF_04": "Cba連接詞", "AF_05": "Cbb關聯連接詞", "AF_06": "D副詞",
        "AF_07": "Da數量副詞", "AF_08": "DE的之得地",
        "AF_09": "Dfa動詞前程度副詞", "AF_10": "Dfb動詞後程度副詞",
        "AF_11": "Di時態標記", "AF_12": "Dk句副詞", "AF_14": "I感嘆詞",
        "AF_15": "Na普通名詞", "AF_16": "Nb專有名詞", "AF_17": "Nc地方詞",
        "AF_18": "Ncd位置詞", "AF_19": "Nd時間詞", "AF_20": "Nep指代定詞",
        "AF_21": "Neqa數量定詞", "AF_22": "Neqb後置數量定詞",
        "AF_23": "Nes特指定詞", "AF_24": "Neu數詞定詞", "AF_25": "Nf量詞",
        "AF_26": "Ng後置詞", "AF_27": "Nh代名詞", "AF_28": "Nv名物化動詞",
        "AF_29": "P介詞", "AF_30": "SHI是", "AF_31": "T語助詞",
        "AF_32": "V2有", "AF_33": "VA動作不及物動詞", "AF_34": "VAC動作使動動詞",
        "AF_35": "VB動作類及物動詞", "AF_36": "VC動作及物動詞",
        "AF_37": "VCL動作接地方賓語動詞", "AF_38": "VD雙賓動詞",
        "AF_39": "VE動作句賓動詞", "AF_40": "VF動作謂賓動詞",
        "AF_41": "VG分類動詞", "AF_42": "VH狀態不及物動詞",
        "AF_43": "VHC狀態使動動詞", "AF_44": "VI狀態類及物動詞",
        "AF_45": "VJ狀態及物動詞", "AF_46": "VK狀態句賓動詞",
        "AF_47": "VL狀態謂賓動詞",
    },
}
INDEX_NAME_MAP = {}
INDEX_CATEGORY_MAP = {}
for _cat, _items in INDEX_CATEGORIES.items():
    for _code, _name in _items.items():
        INDEX_NAME_MAP[_code] = _name
        INDEX_CATEGORY_MAP[_code] = _cat

# ── LLM 報告用情境模組（同 prog04）─────────────────────────────────────────
SCENARIO_MODULE_1 = """## 二、情境特定分析：歷時變化分析

在陳述差異時，必須將「時間順序」作為組織軸線。
請依照以下子項目分析：

### 2.1 語言現代化的跡象
- 詞彙難易度（AC_01, AC_02）是否隨時間呈現規律變化？
- 成語比率（AG_06）在不同時期譯本中的變化方向？
- 若出現非單調變化，明確標示此為反常模式，需進一步解釋。

### 2.2 句式與篇章結構的演變
- 段落數、句數（AA_05, AA_06）的變化是否反映閱讀習慣的改變？
- 詞性分布中，連接詞、代詞比例的變化是否呈現顯著化或隱化趨勢？

### 2.3 譯本間的親疏關係
- 根據分群結果，時序相鄰的譯本是否聚為一群？
- 若出現時序相隔較遠的譯本反而更相似的情況，標記為需進一步探討的現象。

### 禁令補充
- 不得將「較新」直接等同於「較現代」或「較好」。"""

SCENARIO_MODULE_2 = """## 二、情境特定分析：人機差異分析

### 2.1 機器翻譯典型特徵的檢視
根據量化數據檢視機器翻譯是否呈現以下已知特徵：
- 詞彙多樣性是否偏低？（MSTTR, RTTR）
- 句長是否較為整齊（變異數較小）？

### 2.2 人工翻譯的差異化特徵
- 人工譯本在哪些指標上顯著偏離機器譯本？

### 禁令補充
- 不得對機器翻譯系統的品質做整體評價。
- 不得預設「人工一定優於機器」或反之。"""

SCENARIO_MODULE_3 = """## 二、情境特定分析：改寫策略分析

【重要限制】：本系統目前未進行原文對齊，所有關於刪減/增補的陳述必須明確標示為「推測」。

### 3.1 規模差異的量化
- 總字數、總詞數（AA_01, AA_03）的差距幅度

### 3.2 改寫策略的語言訊號
- 編譯本的詞彙密度是否反而較高？（可能反映資訊壓縮策略）
- 編譯本的平均句長是否較短？（可能反映簡化策略）

### 禁令補充
- 嚴禁使用「忠實度」、「完整性」等規範性詞彙。"""

SCENARIO_MODULE_4 = """## 二、情境特定分析：個人風格差異分析

【前提】：所有譯本產出於相近時期，時代因素已被控制。

### 2.1 核心風格指標的差異
- 在哪些指標上，譯者間差異達到統計顯著？

### 2.2 風格穩定性的檢視
- 譯者內變異 vs 譯者間變異的比較

### 禁令補充
- 不得將個人風格與「翻譯能力」掛鉤。"""

SCENARIO_MODULES = {
    "scenario_1": SCENARIO_MODULE_1,
    "scenario_2": SCENARIO_MODULE_2,
    "scenario_3": SCENARIO_MODULE_3,
    "scenario_4": SCENARIO_MODULE_4,
}

SYSTEM_PROMPT_BASE = """# 角色定義
你是一位翻譯研究的量化分析助理。你的工作是將統計數據轉化為結構化的觀察與假設，供研究者進一步驗證。

# 嚴格禁令
1. 僅根據本次提供的量化數據進行分析，不得引入外部知識。
2. 不得使用「優秀」、「忠實」、「準確」、「流暢」等評價性詞彙。
3. 不得推測譯者的主觀意圖，除非數據直接支持且標明為推論。
4. 遇到數據中未涵蓋的問題，明確回答「本次數據無法回答此問題」。
5. 所有統計顯著性的陳述必須基於提供的 p 值或效果量。

# 輸入資料結構
【研究情境】：{scenario_type}
【譯本 metadata】：{metadata}
【量化分析結果】：（見 user message 中的 JSON 資料）
【使用者補充關注點】：{user_focus}

# 通用輸出結構
輸出必須包含以下四個區塊，缺一不可：

## 一、數據事實描述
僅陳述統計結果，每項陳述後標註對應的指標代碼與數值。
區分：達統計顯著的差異（標明 p 值與效果量）、未達顯著但數值差距較大的趨勢、無明顯差異的指標。

{scenario_module}

## 三、待驗證事項
- 需研究者回到原文比對的具體項目
- 數據中的異常值或可疑模式
- 本分析無法處理的面向

## 四、本次分析的已知限制
明確列出：分詞系統潛在誤差、樣本量對統計檢定的影響、指標本身的理論侷限。"""

# ── matplotlib 中文字型 ──────────────────────────────────────────────────────
def _find_cjk_font():
    candidates = ["Microsoft JhengHei","Microsoft YaHei","SimHei",
                  "Noto Sans CJK TC","Noto Sans CJK SC"]
    available = {f.name for f in fm.fontManager.ttflist}
    for c in candidates:
        if c in available:
            return c
    return None

_cjk = _find_cjk_font()
if _cjk:
    plt.rcParams['font.family'] = _cjk
plt.rcParams['axes.unicode_minus'] = False

# ============================================================================
# SECTION 2: Resource Loading (load once at startup)
# ============================================================================

_RESOURCES = {}

def _load_resources():
    global _RESOURCES
    print("【系統啟動】載入參考語料資源...")
    def _read(fname):
        p = CORPUS_INFO / fname
        if not p.exists():
            print(f"  警告：找不到 {fname}")
            return []
        with open(p, 'r', encoding='utf-8') as f:
            return [l.replace('\n','') for l in f.readlines()]

    _RESOURCES['T01']    = _read('T01_成語5335.txt')
    _RESOURCES['T01_B']  = _read('T01_成語20389.txt')
    _RESOURCES['T02']    = _read('T02_國小常用詞.txt')
    _RESOURCES['T03']    = _read('T03_國小常用詞3000.txt')
    _RESOURCES['T04']    = _read('T04_現代漢語3000.txt')
    _RESOURCES['T05']    = _read('T05_現代漢語8000.txt')
    _RESOURCES['T06_V500']  = _read('T06_Sinica高頻動詞前500.txt')
    _RESOURCES['T06_N3000'] = _read('T06_Sinica高頻詞前3000詞_常用名詞.txt')
    _RESOURCES['T06_A3000'] = _read('T06_Sinica高頻詞前3000詞_常用形容詞.txt')
    _RESOURCES['T06_D3000'] = _read('T06_Sinica高頻詞前3000詞_常用副詞.txt')
    _RESOURCES['T06_N8000'] = _read('T06_Sinica高頻詞前8000詞_常用名詞.txt')
    _RESOURCES['T06_A8000'] = _read('T06_Sinica高頻詞前8000詞_常用形容詞.txt')
    _RESOURCES['T06_D8000'] = _read('T06_Sinica高頻詞前8000詞_常用副詞.txt')

    strokes_path = CORPUS_INFO / 'strokesDictionary.json'
    if strokes_path.exists():
        with open(strokes_path, 'r', encoding='utf-8') as f:
            _RESOURCES['strokes'] = json.load(f)
    else:
        _RESOURCES['strokes'] = {}
        print("  警告：找不到 strokesDictionary.json")

    print("【系統啟動】資源載入完成。")

_load_resources()

# ============================================================================
# SECTION 3: LLM API Wrappers
# ============================================================================

def _call_openrouter(api_key, model, system_prompt, user_prompt, max_tokens=16000, temperature=0.5):
    """通用 OpenRouter API 呼叫（同步）。"""
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://retranslation-agent.local",
        "X-Title": "Retranslation Analysis Agent",
    }
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
        "max_tokens": max_tokens,
        "temperature": temperature,
    }
    data = json.dumps(payload).encode("utf-8")
    req  = urllib.request.Request(url, data=data, headers=headers, method="POST")
    try:
        with urllib.request.urlopen(req, timeout=180) as resp:
            result = json.loads(resp.read().decode("utf-8"))
            return result["choices"][0]["message"]["content"]
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8") if e.fp else ""
        raise RuntimeError(f"OpenRouter API HTTP 錯誤 ({e.code}): {body}")
    except Exception as e:
        raise RuntimeError(f"OpenRouter API 呼叫失敗: {e}")


def call_main_agent(system_prompt, user_prompt):
    """主 Agent (LLM1) 呼叫。"""
    return _call_openrouter(LLM1_API_KEY, LLM1_MODEL, system_prompt, user_prompt)


def call_worker_agent(system_prompt, user_prompt):
    """工作程序 Agent (LLM2) 呼叫。"""
    return _call_openrouter(LLM2_API_KEY, LLM2_MODEL, system_prompt, user_prompt)


# ============================================================================
# SECTION 4: Text Format Validation
# ============================================================================

def validate_ckip_format(filepath: Path):
    """
    檢查檔案是否為 ckip_segmentation_sumin 工具產生的斷詞格式。
    判斷依據：
      1. 非空行中，含底線（_）的 token 佔比 > 70%
      2. 至少出現一個 CKIP 特徵標記（CATEGORY 系列或標準詞性碼）
    返回 (is_valid: bool, reason: str)
    """
    CKIP_SIGNATURES = re.compile(
        r'(PERIODCATEGORY|COMMACATEGORY|PAUSECATEGORY|SEMICOLONCATEGORY|'
        r'QUESTIONCATEGORY|EXCLAMATIONCATEGORY|PARENTHESISCATEGORY|'
        r'DASHCATEGORY|ETCCATEGORY|COLONCATEGORY|SPCHANGECATEGORY|'
        r'\b(?:Na|Nb|Nc|Ncd|Nd|Nf|Nh|Nv|VA|VB|VC|VD|VE|VF|VG|VH|VHC|'
        r'VI|VJ|VK|VL|VAC|VCL|DE|Di|SHI|FW|Caa|Cab|Cba|Cbb|'
        r'Neu|Nes|Nep|Neqa|Neqb|Ng|P|D|Da|Dfa|Dfb|Dk|I|T|A)\b)'
    )
    try:
        with open(filepath, 'r', encoding='utf-8-sig') as f:
            lines = f.readlines()
    except Exception as e:
        return False, f"無法讀取檔案：{e}"

    total_tokens = 0
    underscore_tokens = 0
    has_ckip_signature = False
    sample_lines = [l.rstrip() for l in lines if l.strip()][:100]

    if not sample_lines:
        return False, "檔案為空。"

    for line in sample_lines:
        tokens = line.split()
        for tok in tokens:
            total_tokens += 1
            if '_' in tok:
                underscore_tokens += 1
            if not has_ckip_signature and CKIP_SIGNATURES.search(tok):
                has_ckip_signature = True

    if total_tokens == 0:
        return False, "檔案無有效內容。"

    ratio = underscore_tokens / total_tokens
    if ratio < 0.70:
        return False, (
            f"token 中含底線（詞_詞性）的比例僅 {ratio:.0%}，"
            "不符合 ckip_segmentation_sumin 的輸出格式。"
        )
    if not has_ckip_signature:
        return False, "未偵測到 CKIP 詞性標記，可能非 ckip_segmentation_sumin 產生的檔案。"

    return True, "格式符合 ckip_segmentation_sumin 輸出。"


# ============================================================================
# SECTION 5: Indicator Calculation (adapted from Prog02)
# ============================================================================

def _text2strokes(tokens, strokes_db):
    words, strokes, missing = [], [], []
    for tok in tokens:
        for char in tok:
            if char in strokes_db:
                words.append(char)
                strokes.append(strokes_db[char]['strokes_total'])
            else:
                missing.append(char)
    return words, strokes, missing


def calculate_text_index(corpus_file_path: Path, resources: dict) -> dict:
    """
    計算單一文本的所有語言學指標。
    改寫自 Prog02_indicators_multi-v3_1150509.py 的 calculateTextIndex()。
    """
    fname = corpus_file_path.name

    with open(corpus_file_path, 'r', encoding='utf-8-sig') as xf:
        text0 = xf.read()

    text0  = text0.replace('﻿', '')
    text1  = text0.split('\n')
    text2  = [xx.split() for xx in text1]
    text3_1, text3_2 = [], []
    for i, row in enumerate(text2):
        t1, t2 = [], []
        for xxi in row:
            parts = xxi.split('_')
            t1.append(parts[0])
            t2.append(parts[1] if len(parts) > 1 else '')
        text3_1.append(t1)
        text3_2.append(t2)

    # ── 標點 / 句子統計 ───────────────────────────────────────────────────────
    com1 = sum(row.count("COMMACATEGORY") for row in text3_2)
    ss1  = sum(row.count("PERIODCATEGORY")     for row in text3_2)
    ss2  = sum(row.count("SEMICOLONCATEGORY")  for row in text3_2)
    ss3  = sum(row.count("QUESTIONCATEGORY")   for row in text3_2)
    ss4  = sum(row.count("EXCLAMATIONCATEGORY")for row in text3_2)
    sentenceTotal = ss1 + ss2 + ss3 + ss4
    if sentenceTotal == 0:
        sentenceTotal = 1   # 防止除以零

    A_index1 = com1 / sentenceTotal

    # ── 展開成一維 ─────────────────────────────────────────────────────────────
    text3_1_long = [t for row in text3_1 for t in row]
    text3_2_long = [t for row in text3_2 for t in row]

    sen_pos_set = ["PERIODCATEGORY","SEMICOLONCATEGORY","QUESTIONCATEGORY","EXCLAMATIONCATEGORY"]
    sec_pos_set = sen_pos_set + ["COMMACATEGORY","COLONCATEGORY","DASHCATEGORY"]

    def _calc_lengths(boundary_set):
        lengths, xlen = [], 0
        for sss in text3_2_long:
            if sss in boundary_set:
                lengths.append(xlen); xlen = 0
            elif 'CATEGORY' not in sss:
                xlen += 1
        return lengths

    sen_length = _calc_lengths(sen_pos_set)
    sec_length = _calc_lengths(sec_pos_set)
    if not sen_length: sen_length = [0]
    if not sec_length: sec_length = [0]

    mean_sentenceLength = np.mean(sen_length)
    var_sentenceLength  = np.var(sen_length)
    mean_sectionLength  = np.mean(sec_length)
    var_sectionLength   = np.var(sec_length)
    sd_sectionLength    = np.sqrt(var_sectionLength)

    # ── 標點符號總數 ───────────────────────────────────────────────────────────
    categoryTotal = sum(1 for t in text3_2_long if 'CATEGORY' in t)

    # ── 詞 / 詞性統計 ───────────────────────────────────────────────────────────
    tokens, tokens_pos, TokenAddPos, uniTokens, uni_TokenAddPos = [], [], [], [], []
    for ii, row in enumerate(text3_1):
        for jj in range(len(row)):
            if 'CATEGORY' not in text3_2[ii][jj]:
                w = text3_1[ii][jj]
                p = text3_2[ii][jj]
                tokens.append(w)
                tokens_pos.append(p)
                TokenAddPos.append(f"{w}-{p}")
                if w not in uniTokens:         uniTokens.append(w)
                if f"{w}-{p}" not in uni_TokenAddPos: uni_TokenAddPos.append(f"{w}-{p}")

    tokensTotal    = len(tokens)
    uniTokensTotal = len(uni_TokenAddPos)
    if tokensTotal == 0: tokensTotal = 1

    wordTotal   = sum(len(w) for w in tokens)
    words_0     = ''.join(tokens)
    uni_words   = set(words_0)
    uni_words_len = len(uni_words)

    # 詞長
    tokens_len       = [len(w) for w in tokens]
    mean_tokens_len  = np.mean(tokens_len) if tokens_len else 0

    # 段落數
    lineTotal = sum(1 for row in text3_1 if row)

    # 基礎比值
    A_index2 = wordTotal   / sentenceTotal
    A_index3 = tokensTotal / sentenceTotal
    A_index4 = categoryTotal / sentenceTotal
    A_index5 = wordTotal   / max(categoryTotal, 1)
    A_index6 = tokensTotal / max(categoryTotal, 1)
    A_index7 = wordTotal   / max(lineTotal, 1)
    A_index8 = tokensTotal / max(lineTotal, 1)

    # ── 詞頻比對 ──────────────────────────────────────────────────────────────
    R  = resources
    elementaryTotal  = sum(1 for w in tokens if w in R.get('T02', []))
    elementaryTotal3 = sum(1 for w in tokens if w in R.get('T03', []))
    elementaryTotal4 = sum(1 for w in tokens if w in R.get('T04', []))
    elementaryTotal5 = sum(1 for w in tokens if w in R.get('T05', []))

    xx = Counter(tokens)
    idiomTotal  = sum(v for k,v in xx.items() if k in R.get('T01',  []))
    idiomTotal2 = sum(v for k,v in xx.items() if k in R.get('T01_B',[]))

    tokens_V   = [tokens[i] for i in range(len(tokens)) if tokens_pos[i] in Vset]
    tokens_N   = [tokens[i] for i in range(len(tokens)) if tokens_pos[i] in Nset]
    tokens_A   = [tokens[i] for i in range(len(tokens)) if tokens_pos[i] in Aset] or ['XX']
    tokens_ADV = [tokens[i] for i in range(len(tokens)) if tokens_pos[i] in ADVset]

    VelementaryTotal500   = sum(1 for w in tokens_V   if w in R.get('T06_V500', []))
    NelementaryTotal3000  = sum(1 for w in tokens_N   if w in R.get('T06_N3000',[]))
    AelementaryTotal3000  = sum(1 for w in tokens_A   if w in R.get('T06_A3000',[]))
    ADVelementaryTotal3000= sum(1 for w in tokens_ADV if w in R.get('T06_D3000',[]))
    NelementaryTotal8000  = sum(1 for w in tokens_N   if w in R.get('T06_N8000',[]))
    AelementaryTotal8000  = sum(1 for w in tokens_A   if w in R.get('T06_A8000',[]))
    ADVelementaryTotal8000= sum(1 for w in tokens_ADV if w in R.get('T06_D8000',[]))

    # ── 詞性計數 ─────────────────────────────────────────────────────────────
    pos_A_fre    = tokens_pos.count("A")
    pos_V_fre    = sum(1 for p in tokens_pos if p in Vset)
    pos_N_fre    = sum(1 for p in tokens_pos if p in Nset)
    pos_Nh_fre2  = sum(1 for p in tokens_pos if p in Nhset)
    pos_ADV_fre  = sum(1 for p in tokens_pos if p in ADVset)
    pos_ASP_fre  = tokens_pos.count("Di")
    pos_C_fre    = sum(1 for p in tokens_pos if p in Cset)
    pos_Ca_fre   = sum(1 for p in tokens_pos if p in Caset)
    pos_Cb_fre   = sum(1 for p in tokens_pos if p in Cbset)
    pos_DET_fre  = sum(1 for p in tokens_pos if p in DETset)
    pos_P_fre    = tokens_pos.count("P")
    pos_POST_fre = sum(1 for p in tokens_pos if p in POSTset)
    pos_T_fre    = sum(1 for p in tokens_pos if p in Tset)
    pos_DE_fre   = tokens_pos.count("DE")
    pos_SHI_fre  = tokens_pos.count("SHI")
    pos_V2_fre   = tokens_pos.count("V2")

    # 個別詞性
    def _pc(p): return tokens_pos.count(p)
    pos_Caa_fre = _pc("Caa"); pos_Cab_fre = _pc("Cab")
    pos_Cba_fre = _pc("Cba"); pos_Cbb_fre = _pc("Cbb")
    pos_D_fre   = _pc("D");   pos_Da_fre  = _pc("Da")
    pos_Dfa_fre = _pc("Dfa"); pos_Dfb_fre = _pc("Dfb")
    pos_Di_fre  = _pc("Di");  pos_Dk_fre  = _pc("Dk")
    pos_I_fre   = _pc("I");   pos_Na_fre  = _pc("Na")
    pos_Nb_fre  = _pc("Nb");  pos_Nc_fre  = _pc("Nc")
    pos_Ncd_fre = _pc("Ncd"); pos_Nd_fre  = _pc("Nd")
    pos_Nep_fre = _pc("Nep"); pos_Neqa_fre= _pc("Neqa")
    pos_Neqb_fre= _pc("Neqb");pos_Nes_fre = _pc("Nes")
    pos_Neu_fre = _pc("Neu"); pos_Nf_fre  = _pc("Nf")
    pos_Ng_fre  = _pc("Ng");  pos_Nh_fre  = _pc("Nh")
    pos_Nv_fre  = _pc("Nv")
    pos_VA_fre  = _pc("VA");  pos_VAC_fre = _pc("VAC")
    pos_VB_fre  = _pc("VB");  pos_VC_fre  = _pc("VC")
    pos_VCL_fre = _pc("VCL"); pos_VD_fre  = _pc("VD")
    pos_VE_fre  = _pc("VE");  pos_VF_fre  = _pc("VF")
    pos_VG_fre  = _pc("VG");  pos_VH_fre  = _pc("VH")
    pos_VHC_fre = _pc("VHC"); pos_VI_fre  = _pc("VI")
    pos_VJ_fre  = _pc("VJ");  pos_VK_fre  = _pc("VK")
    pos_VL_fre  = _pc("VL")

    # 內容詞
    cont1 = pos_N_fre + pos_A_fre + pos_V_fre
    cont2 = cont1 + pos_ADV_fre

    # AJ 指標
    AJ04_set = ['什麼','哪','什','甚麼','啥']
    AJ05_set = ['它','牠','它們','牠們','其','之','前者','後者','兩者','二者',
                '以上','以下','誰','誰人','何者','什麼','甚麼','啥']
    AJ06_set = ['它','牠','它們','牠們','其','之','前者','後者','兩者','二者','以上','以下']
    AJ07_set = ['我','我們','我方','我倆','咱們','咱','俺','儂','咱家','咱們','予','余',
                '老子','我等','吾等']
    AJ08_set = ['你','妳','你們','你方','妳們','妳方','你倆','妳倆']
    AJ09_set = ['他','她','他們','她們','他倆']
    AJ10_set = ['對方','個人','人家','各位','各自','自己','自我','本人','自身',
                '別人','彼此','雙方','大家','大夥兒','大眾','他人','私人','彼方']
    AJ02_set = ['Caa','Cbb','Cab','Cba']
    AJ03_set = ['Cbb','Cba']

    pos_Nep_AJ04_fre = sum(1 for i in range(len(tokens))
                           if tokens_pos[i]=="Nep" and tokens[i] not in AJ04_set)
    pos_Nh_AJ05_fre  = sum(1 for i in range(len(tokens))
                           if tokens_pos[i]=="Nh"  and tokens[i] not in AJ05_set)
    pos_Nh_AJ06_fre  = sum(1 for i in range(len(tokens))
                           if tokens_pos[i]=="Nh"  and tokens[i] in  AJ06_set)
    pos_Nh_AJ07_fre  = sum(1 for i in range(len(tokens))
                           if tokens_pos[i]=="Nh"  and tokens[i] in  AJ07_set)
    pos_Nh_AJ08_fre  = sum(1 for i in range(len(tokens))
                           if tokens_pos[i]=="Nh"  and tokens[i] in  AJ08_set)
    pos_Nh_AJ09_fre  = sum(1 for i in range(len(tokens))
                           if tokens_pos[i]=="Nh"  and tokens[i] in  AJ09_set)
    pos_Nh_AJ10_fre  = sum(1 for i in range(len(tokens))
                           if tokens_pos[i]=="Nh"  and tokens[i] in  AJ10_set)
    AJ02_fre = sum(1 for p in tokens_pos if p in AJ02_set)
    AJ03_fre = sum(1 for p in tokens_pos if p in AJ03_set)

    # ── 豐富度指標 ─────────────────────────────────────────────────────────────
    TTR = uniTokensTotal / tokensTotal
    perK = 1000
    kk   = len(TokenAddPos) // perK
    kk_TTR = []
    for ii in range(1, kk+1):
        seg   = TokenAddPos[(ii-1)*1000 : ii*1000]
        kk_TTR.append(len(set(seg)) / perK)
    STD_TTR   = sum(kk_TTR) / len(kk_TTR) if kk_TTR else 0
    Guiraud_R = uniTokensTotal / pow(tokensTotal, 0.5)
    LexicalDensity1 = cont1 / tokensTotal
    LexicalDensity2 = cont2 / tokensTotal

    return {
        "Text":  [fname],
        "AA_01": [wordTotal],
        "AA_02": [uni_words_len],
        "AA_03": [tokensTotal],
        "AA_04": [uniTokensTotal],
        "AA_05": [lineTotal],
        "AA_06": [sentenceTotal],
        "AC_01": [round((tokensTotal - elementaryTotal4) / tokensTotal * 1000, 4)],
        "AC_02": [round((tokensTotal - elementaryTotal5) / tokensTotal * 1000, 4)],
        "AD_02": [round(STD_TTR, 4)],
        "AD_03": [round(Guiraud_R, 4)],
        "AG_01": [round(LexicalDensity1, 4)],
        "AG_02": [round(LexicalDensity2, 4)],
        "AG_05": [round(idiomTotal  / tokensTotal * 1000, 4)],
        "AG_06": [round(idiomTotal2 / tokensTotal * 1000, 4)],
        "AH_01": [round(mean_sentenceLength, 4)],
        "AH_03": [round(np.sqrt(var_sentenceLength), 4)],
        "AH_06": [round(mean_sectionLength, 4)],
        "AH_08": [round(sd_sectionLength, 4)],
        "AJ_01": [round(pos_C_fre   / tokensTotal * 1000, 4)],
        "AJ_02": [round(AJ02_fre    / tokensTotal * 1000, 4)],
        "AJ_03": [round(AJ03_fre    / tokensTotal * 1000, 4)],
        "AJ_04": [round(pos_Nep_AJ04_fre / tokensTotal * 1000, 4)],
        "AJ_05": [round(pos_Nh_AJ05_fre  / tokensTotal * 1000, 4)],
        "AJ_06": [round(pos_Nh_AJ06_fre  / tokensTotal * 1000, 4)],
        "AJ_07": [round(pos_Nh_AJ07_fre  / tokensTotal * 1000, 4)],
        "AJ_08": [round(pos_Nh_AJ08_fre  / tokensTotal * 1000, 4)],
        "AJ_09": [round(pos_Nh_AJ09_fre  / tokensTotal * 1000, 4)],
        "AJ_10": [round(pos_Nh_AJ10_fre  / tokensTotal * 1000, 4)],
        "AJ_11": [round(pos_Nh_AJ10_fre  / tokensTotal * 1000, 4)],
        "AF_01": [round(pos_A_fre   / tokensTotal * 1000, 4)],
        "AF_02": [round(pos_Caa_fre / tokensTotal * 1000, 4)],
        "AF_03": [round(pos_Cab_fre / tokensTotal * 1000, 4)],
        "AF_04": [round(pos_Cba_fre / tokensTotal * 1000, 4)],
        "AF_05": [round(pos_Cbb_fre / tokensTotal * 1000, 4)],
        "AF_06": [round(pos_D_fre   / tokensTotal * 1000, 4)],
        "AF_07": [round(pos_Da_fre  / tokensTotal * 1000, 4)],
        "AF_08": [round(pos_DE_fre  / tokensTotal * 1000, 4)],
        "AF_09": [round(pos_Dfa_fre / tokensTotal * 1000, 4)],
        "AF_10": [round(pos_Dfb_fre / tokensTotal * 1000, 4)],
        "AF_11": [round(pos_Di_fre  / tokensTotal * 1000, 4)],
        "AF_12": [round(pos_Dk_fre  / tokensTotal * 1000, 4)],
        "AF_14": [round(pos_I_fre   / tokensTotal * 1000, 4)],
        "AF_15": [round(pos_Na_fre  / tokensTotal * 1000, 4)],
        "AF_16": [round(pos_Nb_fre  / tokensTotal * 1000, 4)],
        "AF_17": [round(pos_Nc_fre  / tokensTotal * 1000, 4)],
        "AF_18": [round(pos_Ncd_fre / tokensTotal * 1000, 4)],
        "AF_19": [round(pos_Nd_fre  / tokensTotal * 1000, 4)],
        "AF_20": [round(pos_Nep_fre / tokensTotal * 1000, 4)],
        "AF_21": [round(pos_Neqa_fre/ tokensTotal * 1000, 4)],
        "AF_22": [round(pos_Neqb_fre/ tokensTotal * 1000, 4)],
        "AF_23": [round(pos_Nes_fre / tokensTotal * 1000, 4)],
        "AF_24": [round(pos_Neu_fre / tokensTotal * 1000, 4)],
        "AF_25": [round(pos_Nf_fre  / tokensTotal * 1000, 4)],
        "AF_26": [round(pos_Ng_fre  / tokensTotal * 1000, 4)],
        "AF_27": [round(pos_Nh_fre  / tokensTotal * 1000, 4)],
        "AF_28": [round(pos_Nv_fre  / tokensTotal * 1000, 4)],
        "AF_29": [round(pos_P_fre   / tokensTotal * 1000, 4)],
        "AF_30": [round(pos_SHI_fre / tokensTotal * 1000, 4)],
        "AF_31": [round(pos_T_fre   / tokensTotal * 1000, 4)],
        "AF_32": [round(pos_V2_fre  / tokensTotal * 1000, 4)],
        "AF_33": [round(pos_VA_fre  / tokensTotal * 1000, 4)],
        "AF_34": [round(pos_VAC_fre / tokensTotal * 1000, 4)],
        "AF_35": [round(pos_VB_fre  / tokensTotal * 1000, 4)],
        "AF_36": [round(pos_VC_fre  / tokensTotal * 1000, 4)],
        "AF_37": [round(pos_VCL_fre / tokensTotal * 1000, 4)],
        "AF_38": [round(pos_VD_fre  / tokensTotal * 1000, 4)],
        "AF_39": [round(pos_VE_fre  / tokensTotal * 1000, 4)],
        "AF_40": [round(pos_VF_fre  / tokensTotal * 1000, 4)],
        "AF_41": [round(pos_VG_fre  / tokensTotal * 1000, 4)],
        "AF_42": [round(pos_VH_fre  / tokensTotal * 1000, 4)],
        "AF_43": [round(pos_VHC_fre / tokensTotal * 1000, 4)],
        "AF_44": [round(pos_VI_fre  / tokensTotal * 1000, 4)],
        "AF_45": [round(pos_VJ_fre  / tokensTotal * 1000, 4)],
        "AF_46": [round(pos_VK_fre  / tokensTotal * 1000, 4)],
        "AF_47": [round(pos_VL_fre  / tokensTotal * 1000, 4)],
    }


def run_indicator_calculation(file_paths: list, output_dir: Path, resources: dict):
    """
    對所有文本計算指標，儲存 xlsx + csv，回傳 (DataFrame, xlsx_path, csv_path)。
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    result_df = pd.DataFrame()
    for fp in file_paths:
        try:
            row_dict = calculate_text_index(Path(fp), resources)
            result_df = pd.concat([result_df, pd.DataFrame(row_dict)],
                                  axis=0, ignore_index=True)
        except Exception as e:
            print(f"  計算 {fp} 時發生錯誤：{e}")
            traceback.print_exc()

    timestamp = now_tw().strftime("%Y%m%d%H%M%S")
    xlsx_path = output_dir / f"textIndex_calculated_1_{timestamp}.xlsx"
    csv_path  = output_dir / "textIndex_calculated.csv"
    result_df.to_excel(str(xlsx_path), index=False)
    result_df.to_csv(str(csv_path),    index=False, encoding='utf-8-sig')
    return result_df, xlsx_path, csv_path


# ============================================================================
# SECTION 6: Statistical Analysis (adapted from prog03, no tkinter)
# ============================================================================

def cohens_d(g1, g2):
    n1, n2 = len(g1), len(g2)
    if n1 < 2 or n2 < 2:
        return 0.0
    var1, var2 = np.var(g1, ddof=1), np.var(g2, ddof=1)
    pooled = np.sqrt(((n1-1)*var1 + (n2-1)*var2) / (n1+n2-2))
    return (np.mean(g1) - np.mean(g2)) / pooled if pooled != 0 else 0.0


def eta_squared(group_data):
    all_data   = np.concatenate(group_data)
    grand_mean = all_data.mean()
    ss_between = sum(len(g) * (np.mean(g) - grand_mean)**2 for g in group_data)
    ss_total   = np.sum((all_data - grand_mean)**2)
    return ss_between / ss_total if ss_total != 0 else 0.0


def effect_label(val, method):
    if method == "t-test":
        v = abs(val)
        return "大" if v>=0.8 else "中" if v>=0.5 else "小" if v>=0.2 else "極小"
    else:
        return "大" if val>=0.14 else "中" if val>=0.06 else "小" if val>=0.01 else "極小"


def run_statistical_analysis(csv_path: Path, output_dir: Path):
    """
    執行完整統計分析（敘述統計、差異分析、隨機森林、分群）。
    回傳 dict 包含所有輸出檔路徑與摘要資訊。
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    # ── 1. 讀入資料 ────────────────────────────────────────────────────────────
    df = pd.read_csv(str(csv_path))

    df['Type1'] = df['Text'].str[0].str.upper()
    df['Type2'] = df['Type1']
    GROUP_COL   = 'Type1'
    N_CLUSTERS  = max(2, len(df['Type1'].unique()))

    # 重新命名欄位（加中文名稱）
    rename_map = {}
    for col in df.columns:
        if col in INDEX_NAME_MAP:
            rename_map[col] = f"{col}_{INDEX_NAME_MAP[col]}"
    df.rename(columns=rename_map, inplace=True)
    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()

    # ── 2. 敘述統計（涵蓋所有計算指標，與語言指標分析一致）────────────────────
    def _desc_row(cat, code, cname, rcol):
        s = df[rcol].dropna()
        if len(s) == 0:
            return None
        return {
            "類別": cat, "指標代碼": code, "指標名稱": cname,
            "文本個數": int(s.count()),
            "平均數":   round(float(s.mean()), 4),
            "中位數":   round(float(s.median()), 4),
            "標準差":   round(float(s.std()), 4),
            "最小值":   round(float(s.min()), 4),
            "最大值":   round(float(s.max()), 4),
            "變異係數(%)": round(float(s.std()/s.mean()*100), 2) if s.mean()!=0 else None,
        }

    all_desc_rows = []
    covered = set()
    # (a) 已分類指標（依分類順序呈現）
    for cat, items in INDEX_CATEGORIES.items():
        for code, cname in items.items():
            rcol = f"{code}_{cname}"
            if rcol not in df.columns:
                continue
            r = _desc_row(cat, code, cname, rcol)
            if r:
                all_desc_rows.append(r)
                covered.add(rcol)
    # (b) 未分類指標（如 AJ 系列），補入確保與指標計算數量一致
    for rcol in numeric_cols:
        if rcol in covered:
            continue
        parts = rcol.split("_")
        code  = (parts[0] + "_" + parts[1]) if len(parts) >= 2 else rcol
        cname = INDEX_NAME_MAP.get(code, code)
        r = _desc_row("其他指標", code, cname, rcol)
        if r:
            all_desc_rows.append(r)
    desc_df   = pd.DataFrame(all_desc_rows)
    desc_path = output_dir / "01_descriptive_stats.csv"
    desc_df.to_csv(str(desc_path), index=False, encoding='utf-8-sig')

    # ── 3. 差異性分析 ──────────────────────────────────────────────────────────
    groups   = df[GROUP_COL].unique()
    n_groups = len(groups)
    method_name = "Independent t-test (Welch)" if n_groups == 2 else "One-way ANOVA"

    diff_results = []
    for col in numeric_cols:
        gdata = [df.loc[df[GROUP_COL]==g, col].dropna().values for g in groups]
        if any(len(g)==0 for g in gdata): continue
        row = {"指標": col}
        for gi, g in enumerate(groups):
            row[f"{g}_N"]  = len(gdata[gi])
            row[f"{g}_M"]  = round(float(np.mean(gdata[gi])), 4)
            row[f"{g}_SD"] = round(float(np.std(gdata[gi], ddof=1)), 4) if len(gdata[gi])>1 else 0.0
        if n_groups == 2:
            stat, p  = stats.ttest_ind(gdata[0], gdata[1], equal_var=False)
            es       = cohens_d(gdata[0], gdata[1])
            es_name  = "Cohen's d"; method = "t-test"
        else:
            stat, p  = stats.f_oneway(*gdata)
            es       = eta_squared(gdata)
            es_name  = "η²"; method = "ANOVA"
        row["方法"]   = method
        row["統計量"] = round(float(stat), 4) if not np.isnan(stat) else None
        row["p值"]    = round(float(p),    4) if not np.isnan(p)    else None
        row["顯著"]   = ("***" if p is not None and p<0.001 else
                         "**"  if p is not None and p<0.01  else
                         "*"   if p is not None and p<0.05  else "ns")
        row["效果量名稱"] = es_name
        row["效果量"]     = round(float(es), 4)
        row["效果大小"]   = effect_label(es, method)
        diff_results.append(row)

    diff_df  = pd.DataFrame(diff_results)
    sig_df   = diff_df[diff_df['顯著'] != "ns"]
    diff_path = output_dir / "02_difference_analysis.csv"
    diff_df.to_csv(str(diff_path), index=False, encoding='utf-8-sig')

    # boxplot（最多 12 個顯著指標）
    box_path = None
    if len(sig_df) > 0:
        sig_cols = sig_df['指標'].tolist()[:12]
        n_row    = max(1, (len(sig_cols) + 2) // 3)
        fig, axes = plt.subplots(n_row, 3, figsize=(14, 4*n_row))
        axes = np.array(axes).flatten()
        for i, col in enumerate(sig_cols):
            ax = axes[i]
            pd_data = [df.loc[df[GROUP_COL]==g, col].dropna().values for g in groups]
            ax.boxplot(pd_data)
            ax.set_xticks(range(1, len(groups) + 1))
            ax.set_xticklabels(list(groups))
            pv = diff_df.loc[diff_df['指標']==col, 'p值'].values
            pv_str = f"p={pv[0]:.4f}" if len(pv)>0 else ""
            ax.set_title(f"{col}\n({pv_str})", fontsize=8)
        for j in range(len(sig_cols), len(axes)):
            axes[j].set_visible(False)
        plt.suptitle(f"Significant Variables by {GROUP_COL}", fontsize=11, y=1.01)
        plt.tight_layout()
        box_path = output_dir / "02_significant_boxplots.png"
        plt.savefig(str(box_path), dpi=150, bbox_inches='tight')
        plt.close()

    # ── 4. 隨機森林 ────────────────────────────────────────────────────────────
    le = LabelEncoder()
    y  = le.fit_transform(df[GROUP_COL])
    X  = df[numeric_cols].fillna(df[numeric_cols].median())

    min_class = pd.Series(y).value_counts().min()
    test_sz   = max(len(np.unique(y)), int(len(df) * 0.2))

    if len(df) > 6 and min_class >= 2 and test_sz < len(df):
        X_train, X_test, y_train, y_test = train_test_split(
            X, y, test_size=test_sz, random_state=42, stratify=y)
        do_eval = True
    else:
        X_train, y_train = X, y
        X_test,  y_test  = X, y
        do_eval = False

    rf = RandomForestClassifier(n_estimators=500, random_state=42)
    rf.fit(X_train, y_train)
    y_pred   = rf.predict(X_test)
    accuracy = accuracy_score(y_test, y_pred)

    imp_df  = pd.DataFrame({"指標": numeric_cols, "重要性": rf.feature_importances_})\
                .sort_values("重要性", ascending=False).reset_index(drop=True)
    imp_path = output_dir / "03_feature_importance.csv"
    imp_df.to_csv(str(imp_path), index=False, encoding='utf-8-sig')

    top_n = min(20, len(imp_df))
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.barh(imp_df['指標'][:top_n][::-1], imp_df['重要性'][:top_n][::-1], color='steelblue')
    ax.set_xlabel("Feature Importance")
    ax.set_title(f"Random Forest — Top {top_n} Feature Importance")
    plt.tight_layout()
    imp_fig_path = output_dir / "03_feature_importance.png"
    plt.savefig(str(imp_fig_path), dpi=150, bbox_inches='tight')
    plt.close()

    # ── 5. 非監督式分群 ─────────────────────────────────────────────────────────
    scaler   = StandardScaler()
    X_scaled = scaler.fit_transform(X)
    pca      = PCA(n_components=min(2, X_scaled.shape[1]), random_state=42)
    X_pca    = pca.fit_transform(X_scaled)

    # 階層式分群
    linkage_mat  = linkage(X_scaled, method='ward')
    fig, ax = plt.subplots(figsize=(10, 5))
    dendrogram(linkage_mat, labels=df['Text'].values, leaf_rotation=45, leaf_font_size=8, ax=ax)
    ax.set_title("Hierarchical Clustering Dendrogram (Ward)")
    ax.set_xlabel("Sample"); ax.set_ylabel("Distance")
    plt.tight_layout()
    dendro_path = output_dir / "04a_dendrogram.png"
    plt.savefig(str(dendro_path), dpi=150, bbox_inches='tight')
    plt.close()

    hc = AgglomerativeClustering(n_clusters=N_CLUSTERS, linkage='ward')
    df['HC_Cluster'] = hc.fit_predict(X_scaled)

    # K-means + Elbow
    max_k   = min(8, max(2, len(df)-1))
    inertia = []
    for k in range(1, max_k+1):
        km = KMeans(n_clusters=k, random_state=42, n_init=10)
        km.fit(X_scaled)
        inertia.append(km.inertia_)
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.plot(list(range(1, max_k+1)), inertia, marker='o', color='coral')
    ax.set_xlabel("Number of Clusters (K)"); ax.set_ylabel("Inertia (SSE)")
    ax.set_title("K-means Elbow Curve"); ax.set_xticks(list(range(1, max_k+1)))
    plt.tight_layout()
    elbow_path = output_dir / "04b_kmeans_elbow.png"
    plt.savefig(str(elbow_path), dpi=150, bbox_inches='tight')
    plt.close()

    km_final = KMeans(n_clusters=N_CLUSTERS, random_state=42, n_init=10)
    df['KM_Cluster'] = km_final.fit_predict(X_scaled)

    # PCA 2D 視覺化
    fig, axes = plt.subplots(1, 2, figsize=(14, 5))
    for ax, lcol, title in zip(axes,
            ['HC_Cluster','KM_Cluster'],
            ['Hierarchical Clustering', f'K-means (K={N_CLUSTERS})']):
        ulabels = sorted(df[lcol].unique())
        colors  = plt.cm.tab10(np.linspace(0, 1, len(ulabels)))
        for idx, cl in enumerate(ulabels):
            mask = df[lcol] == cl
            ax.scatter(X_pca[mask,0], X_pca[mask,1],
                       c=[colors[idx]], label=f"Cluster {cl}", s=100, edgecolors='k')
        for i, txt in enumerate(df['Text']):
            ax.annotate(txt[:8], (X_pca[i,0], X_pca[i,1]),
                        fontsize=7, textcoords='offset points', xytext=(5,3))
        ax.set_xlabel(f"PC1 ({pca.explained_variance_ratio_[0]:.1%})" if X_scaled.shape[1]>=2 else "PC1")
        ax.set_ylabel(f"PC2 ({pca.explained_variance_ratio_[1]:.1%})" if X_scaled.shape[1]>=2 else "PC2")
        ax.set_title(title); ax.legend()
    plt.suptitle("Clustering Results (PCA 2D Projection)", fontsize=12)
    plt.tight_layout()
    cluster_fig_path = output_dir / "04c_clustering_pca2d.png"
    plt.savefig(str(cluster_fig_path), dpi=150, bbox_inches='tight')
    plt.close()

    # 分群結果儲存
    result_cols = ['Text','Type1','Type2',GROUP_COL,'HC_Cluster','KM_Cluster']
    result_cols = [c for c in result_cols if c in df.columns]
    cluster_csv_path = output_dir / "05_cluster_results.csv"
    df[result_cols].to_csv(str(cluster_csv_path), index=False, encoding='utf-8-sig')

    # ── 6. 產出 docx 技術報告 ──────────────────────────────────────────────────
    timestamp   = now_tw().strftime("%Y%m%d_%H%M%S")
    report_path = output_dir / f"Final_Report-{timestamp}.docx"
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微軟正黑體'
    style.font.size = Pt(10)

    title = doc.add_heading('重譯分析技術報告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'報告產出時間：{now_tw().strftime("%Y-%m-%d %H:%M:%S")}（台灣時間）',
                      style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'分析文本數：{df.shape[0]} 筆',
                      style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'分組方式：{GROUP_COL}',
                      style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    # 敘述統計
    doc.add_heading('一、敘述性統計', level=1)
    TABLE_HEADERS = ["指標代碼","指標名稱","文本個數","平均數","中位數","標準差","最小值","最大值","變異係數(%)"]
    for cat, items in INDEX_CATEGORIES.items():
        cat_data = desc_df[desc_df['類別'] == cat]
        if cat_data.empty: continue
        doc.add_heading(f'（{cat}）', level=2)
        tbl = doc.add_table(rows=1+len(cat_data), cols=len(TABLE_HEADERS))
        tbl.style = 'Light Shading Accent 1'
        for j, h in enumerate(TABLE_HEADERS):
            cell = tbl.rows[0].cells[j]
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.bold = True; r.font.size = Pt(8)
        for i, (_, row) in enumerate(cat_data.iterrows(), start=1):
            vals = [str(row['指標代碼']), str(row['指標名稱']),
                    str(int(row['文本個數'])),
                    f"{row['平均數']:.4f}", f"{row['中位數']:.4f}",
                    f"{row['標準差']:.4f}", f"{row['最小值']:.4f}",
                    f"{row['最大值']:.4f}",
                    f"{row['變異係數(%)']:.2f}" if row['變異係數(%)'] is not None else "-"]
            for j, val in enumerate(vals):
                cell = tbl.rows[i].cells[j]
                cell.text = val
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs: r.font.size = Pt(8)
        doc.add_paragraph('')

    # 差異性分析
    doc.add_page_break()
    doc.add_heading('二、差異性分析', level=1)
    doc.add_paragraph(f'分組依據：{GROUP_COL}（共 {n_groups} 組）。採用 {method_name}，顯著水準 α = {ALPHA}。')
    doc.add_paragraph(f'共 {len(diff_df)} 個指標，其中 {len(sig_df)} 個達顯著差異（p < {ALPHA}）。')
    group_sub_headers = []
    for g in groups:
        group_sub_headers.extend([f"{g}\nN", f"{g}\nM", f"{g}\nSD"])
    diff_headers = ["指標"] + group_sub_headers + ["方法","統計量","p值","顯著","效果量","效果大小"]
    diff_tbl = doc.add_table(rows=1+len(diff_df), cols=len(diff_headers))
    diff_tbl.style = 'Light Shading Accent 1'
    for j, h in enumerate(diff_headers):
        cell = diff_tbl.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(7)
    for i, (_, row) in enumerate(diff_df.iterrows(), start=1):
        vals = [str(row['指標'])]
        for g in groups:
            vals.extend([str(int(row[f"{g}_N"])),
                         f"{row[f'{g}_M']:.4f}",
                         f"{row[f'{g}_SD']:.4f}"])
        vals.extend([str(row['方法']),
                     f"{row['統計量']:.4f}" if row['統計量'] is not None else "-",
                     f"{row['p值']:.4f}"    if row['p值']    is not None else "-",
                     str(row['顯著']), f"{row['效果量']:.4f}", str(row['效果大小'])])
        for j, val in enumerate(vals):
            cell = diff_tbl.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(7)
                    if row['顯著'] != "ns" and j == len(vals)-4:
                        r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    doc.add_paragraph('')
    if box_path and box_path.exists():
        doc.add_heading('顯著指標箱形圖', level=2)
        doc.add_picture(str(box_path), width=Cm(16))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 隨機森林
    doc.add_page_break()
    doc.add_heading('三、監督式學習 — 隨機森林特徵重要性', level=1)
    doc.add_paragraph(f'模型準確率（Accuracy）：{accuracy:.4f}')
    if imp_fig_path.exists():
        doc.add_picture(str(imp_fig_path), width=Cm(15))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 分群
    doc.add_page_break()
    doc.add_heading('四、非監督式學習 — 集群分析', level=1)
    if dendro_path.exists():
        doc.add_heading('樹狀圖（階層式分群）', level=2)
        doc.add_picture(str(dendro_path), width=Cm(15))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if elbow_path.exists():
        doc.add_heading('Elbow Curve（K-means）', level=2)
        doc.add_picture(str(elbow_path), width=Cm(12))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cluster_fig_path.exists():
        doc.add_heading('分群視覺化（PCA 2D）', level=2)
        doc.add_picture(str(cluster_fig_path), width=Cm(16))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(report_path))

    return {
        "df":               df,
        "desc_path":        desc_path,
        "diff_path":        diff_path,
        "imp_path":         imp_path,
        "imp_fig_path":     imp_fig_path,
        "box_path":         box_path,
        "dendro_path":      dendro_path,
        "elbow_path":       elbow_path,
        "cluster_fig_path": cluster_fig_path,
        "cluster_csv_path": cluster_csv_path,
        "report_path":      report_path,
        "sig_count":        len(sig_df),
        "total_indicators": len(diff_df),
        "accuracy":         accuracy,
        "groups":           list(groups),
        "n_groups":         n_groups,
        "method_name":      method_name,
        "GROUP_COL":        GROUP_COL,
        "pca_ev":           pca.explained_variance_ratio_.tolist(),
    }


# ============================================================================
# SECTION 7: LLM Report Writing (adapted from prog04, uses LLM1)
# ============================================================================

def load_analysis_data_for_llm(csv_path: Path, output_dir: Path) -> dict:
    """讀取分析資料組成 JSON 結構。"""
    data = {}
    if csv_path.exists():
        raw = pd.read_csv(str(csv_path))
        aj_cols = [c for c in raw.columns if c.startswith('AJ_')]
        raw.drop(columns=aj_cols, inplace=True, errors='ignore')
        data["texts"]      = raw['Text'].tolist()
        data["text_count"] = len(raw)
    else:
        data["texts"] = []; data["text_count"] = 0

    for fname, key in [
        ("01_descriptive_stats.csv",   "descriptive_stats"),
        ("02_difference_analysis.csv", "difference_analysis"),
        ("03_feature_importance.csv",  "feature_importance_top20"),
        ("05_cluster_results.csv",     "cluster_results"),
    ]:
        p = output_dir / fname
        if p.exists():
            d = pd.read_csv(str(p))
            if key == "feature_importance_top20":
                data[key] = d.head(20).to_dict(orient='records')
            else:
                data[key] = d.to_dict(orient='records')
            if key == "difference_analysis":
                sig = d[d['顯著'] != 'ns'] if '顯著' in d.columns else pd.DataFrame()
                data["significant_indicators"] = sig.to_dict(orient='records')
                data["significant_count"]      = len(sig)
                data["total_indicators"]       = len(d)
        else:
            data[key] = []
    if "significant_count" not in data:
        data["significant_count"] = 0
    if "total_indicators"   not in data:
        data["total_indicators"]  = 0
    return data


def run_report_writing(scenario_key: str, user_metadata: str, user_focus: str,
                       csv_path: Path, output_dir: Path) -> dict:
    """
    呼叫主 Agent (LLM1) 撰寫研究報告，儲存 docx + txt，回傳路徑 dict。
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    analysis_data = load_analysis_data_for_llm(csv_path, output_dir)

    scenario_info   = SCENARIOS.get(scenario_key, SCENARIOS["1"])
    scenario_mod_key = scenario_info.get("scenario_module_key", "scenario_1")
    scenario_module  = SCENARIO_MODULES.get(scenario_mod_key, "")

    metadata_str = f"共 {analysis_data['text_count']} 個譯本\n"
    for i, t in enumerate(analysis_data["texts"], 1):
        metadata_str += f"  譯本{i}：{t}\n"
    if user_metadata:
        metadata_str += f"\n【使用者補充】：\n{user_metadata}"

    system_prompt = SYSTEM_PROMPT_BASE.format(
        scenario_type   = scenario_info["title"],
        metadata        = metadata_str,
        user_focus      = user_focus if user_focus else "無",
        scenario_module = scenario_module,
    )
    user_prompt = (
        f"請根據以下量化分析結果，按照 system prompt 中的輸出結構進行深度分析。\n\n"
        f"【量化分析資料（JSON）】：\n```json\n"
        f"{json.dumps(analysis_data, ensure_ascii=False, indent=2)}\n```\n\n"
        "請嚴格按照四個區塊（一、數據事實描述 / 二、情境特定分析 / "
        "三、待驗證事項 / 四、已知限制）輸出分析報告。"
    )

    result_text = call_main_agent(system_prompt, user_prompt)

    timestamp    = now_tw().strftime("%Y%m%d_%H%M%S")
    docx_path    = output_dir / f"LLM_Analysis-{timestamp}.docx"
    txt_path     = output_dir / f"LLM_Analysis-{timestamp}.txt"

    # 存 docx
    doc   = Document()
    style = doc.styles['Normal']
    style.font.name = '微軟正黑體'
    style.font.size = Pt(11)
    heading = doc.add_heading('LLM 深度重譯研究報告', level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'研究情境：{scenario_info["title"]}',
                      style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'報告時間：{now_tw().strftime("%Y-%m-%d %H:%M:%S")}（台灣時間）',
                      style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    for line in result_text.split('\n'):
        stripped = line.strip()
        if   stripped.startswith('#### '):  doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith('### '):   doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):    doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):     doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('- '):     doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith('* '):     doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped:                      doc.add_paragraph(stripped)
    doc.save(str(docx_path))

    # 存 txt
    with open(str(txt_path), 'w', encoding='utf-8') as f:
        f.write(f"研究情境：{scenario_info['title']}\n")
        f.write(f"時間：{now_tw().strftime('%Y-%m-%d %H:%M:%S')}（台灣時間）\n")
        f.write("=" * 60 + "\n\n")
        f.write(result_text)

    return {
        "docx_path":     docx_path,
        "txt_path":      txt_path,
        "result_text":   result_text,
    }


# ============================================================================
# SECTION 8: (Reserved)
# ============================================================================


# ============================================================================
# SECTION 9: Conversation Flow & State Machine
# ============================================================================

STATES = {
    "SCENARIO_SELECT": "研究情境選擇",
    "UPLOAD_GUIDE":    "文本上傳說明",
    "UPLOAD_WAIT":     "等待文本上傳",
    "FORMAT_CHECK":    "格式檢查中",
    "CALCULATING":     "指標計算中",
    "ANALYZING":       "統計分析中",
    "REPORT_WRITING":  "研究報告撰寫中",
    "COMPLETE":        "分析完成",
}

def _assistant_msg(text: str):
    return {"role": "assistant", "content": text}

def _user_msg(text: str):
    return {"role": "user", "content": text}


def make_initial_state(session_id: str) -> dict:
    sess_dir = SESSIONS_DIR / session_id
    sess_dir.mkdir(parents=True, exist_ok=True)
    (sess_dir / "uploads").mkdir(exist_ok=True)
    (sess_dir / "output").mkdir(exist_ok=True)
    return {
        "session_id":     session_id,
        "session_dir":    str(sess_dir),
        "state":          "SCENARIO_SELECT",
        "scenario":       None,
        "uploaded_files": [],
        "invalid_info":   {},
        "cleaned_files":  [],
        "output_files":   [],
        "zip_path":       None,
        "user_metadata":  "",
        "user_focus":     "",
        "analysis_result": None,
        "done_steps":     [],    # 已完成步驟（1-based）
        "current_step":   None,  # 目前執行步驟
    }


def get_intro_message() -> str:
    sys_p = (
        "你是「重譯文本分析系統」的主 Agent（智能助理），"
        "請用繁體中文、專業而友善的語氣向使用者介紹本系統，"
        "並說明四種研究情境，最後邀請使用者選擇情境。"
        "回覆請控制在 400 字以內。"
    )
    user_p = (
        "請介紹本系統並列出以下四種研究情境，請使用者選擇：\n"
        "1. 同一原文的人工翻譯 vs 機器翻譯\n"
        "2. 全譯 vs 編譯（節譯、改寫）\n"
        "3. 不同譯者的同期譯本\n"
        "4. 同一原文的不同時期人工譯本\n"
        "詢問使用者：「請輸入數字 1、2、3 或 4 選擇您要進行的研究情境。」"
    )
    try:
        return call_main_agent(sys_p, user_p)
    except Exception as e:
        return (
            "歡迎使用 TransPrism II 分析引擎 SpectraEngine。\n\n"
            "請依照您的研究設計選擇分析情境：\n\n"
            "1 人工翻譯 vs 機器翻譯\n"
            "2 全譯 vs 編譯／節譯／改寫\n"
            "3 同期不同譯者譯本比較\n"
            "4 不同時期重譯本的歷時比較\n\n"
            "輸入 1、2、3 或 4 。"
        )


def _build_todo_list(sc_num: str, done_steps: list = None, current_step: int = None) -> str:
    """
    即時 to-do list 執行進度（確定性，不依賴 LLM）。
    done_steps: 已完成步驟（1-based list）
    current_step: 目前執行步驟（1-based int）
    """
    done_steps = done_steps or []
    sc_info = SCENARIOS.get(sc_num, {})
    sc_title = sc_info.get("short", sc_info.get("title", ""))

    steps = [
        (1, "情境選擇"),
        (2, "上傳引導"),
        (3, "文本上傳"),
        (4, "格式驗證"),
        (5, "指標計算"),
        (6, "統計分析"),
        (7, "報告撰寫"),
        (8, "成果下載"),
    ]

    total = len(steps)
    done_count = len([s for s in done_steps if s <= total])
    pct = int(done_count / total * 100)

    lines = [
        f"**執行進度｜{sc_title}**",
        f"`{'█' * (done_count * 2)}{'░' * ((total - done_count) * 2)}` {pct}%",
        "",
    ]
    for num, name in steps:
        if num in done_steps:
            lines.append(f"✅ ~~{num}. {name}~~")
        elif num == current_step:
            lines.append(f"⏳ **{num}. {name}**")
        else:
            lines.append(f"☐ {num}. {name}")

    return "\n\n".join(lines)


def handle_scenario_select(user_input: str, state: dict):
    """處理研究情境選擇輸入，回傳 (reply_text, new_state)。"""
    choice = user_input.strip()
    m = re.search(r'[1-4]', choice)
    if not m:
        return ("請輸入數字 **1、2、3 或 4** 選擇研究情境。", state)

    sc_num = m.group()
    sc_info = SCENARIOS[sc_num]
    state["scenario"] = sc_num
    state["done_steps"] = [1]
    state["current_step"] = 2

    sys_p = (
        "你是重譯文本分析系統的主 Agent，請用繁體中文確認使用者的研究情境選擇，"
        "並說明接下來將由工作程序 Agent 協助上傳文本。語氣專業友善，50 字以內。"
    )
    user_p = f"使用者選擇了研究情境 {sc_num}：{sc_info['title']}"
    try:
        confirm_text = call_main_agent(sys_p, user_p)
    except Exception:
        confirm_text = f"已記錄研究情境：**{sc_info['title']}**。"

    state["state"] = "UPLOAD_GUIDE"
    return (confirm_text, state)


def get_upload_guide_message(scenario_num: str) -> str:
    """工作程序 Agent 根據情境提供上傳說明。"""
    sc_info = SCENARIOS[scenario_num]
    sys_p = (
        "你是重譯文本分析系統的工作程序 Agent，請用繁體中文說明文本上傳格式要求。"
        "語氣清晰、條列式，120 字以內。"
        "必須強調：文本必須使用 https://github.com/suhsiung/ckip_segmentation_sumin 工具完成斷詞，"
        "才能上傳。最後提示使用者上傳檔案。"
    )
    user_p = f"研究情境：{sc_info['title']}\n上傳說明：\n{sc_info['upload_prompt']}"
    try:
        return call_worker_agent(sys_p, user_p)
    except:
        return (
            f"**【工作程序 Agent】文本上傳說明**\n\n"
            f"{sc_info['upload_prompt']}\n\n"
            "請在下方選擇並上傳文本檔案（可多選）。"
        )


def handle_uploaded_files(files, state: dict):
    """
    接收 Gradio 上傳的檔案，複製到 session 目錄，回傳狀態更新與訊息。
    """
    if not files:
        return "尚未偵測到上傳檔案，請重新上傳。", state

    uploads_dir = Path(state["session_dir"]) / "uploads"
    uploads_dir.mkdir(parents=True, exist_ok=True)

    saved_paths = []
    for f in files:
        src = Path(f.name) if hasattr(f, 'name') else Path(str(f))
        dst = uploads_dir / src.name
        shutil.copy2(str(src), str(dst))
        saved_paths.append(str(dst))

    state["uploaded_files"] = saved_paths
    state["state"] = "FORMAT_CHECK"

    # 工作程序 Agent 說明開始檢查
    file_list_str = "\n".join([Path(p).name for p in saved_paths])
    sys_p = (
        "你是重譯文本分析系統的工作程序 Agent，"
        "請用繁體中文簡短告知使用者已收到檔案並正在進行格式檢查，50 字以內。"
    )
    user_p = f"已接收以下 {len(saved_paths)} 個檔案：\n{file_list_str}"
    try:
        msg = call_worker_agent(sys_p, user_p)
    except:
        msg = f"已接收 **{len(saved_paths)}** 個檔案，正在進行格式檢查……"

    return msg, state


def run_format_check(state: dict):
    """
    檢查所有上傳檔案是否為 ckip_segmentation_sumin 產生的格式。
    通過 → CALCULATING；不通過 → UPLOAD_WAIT（要求重新上傳）。
    返回 (reply_text, new_state, has_errors)
    """
    failed = {}
    for fp in state["uploaded_files"]:
        ok, reason = validate_ckip_format(Path(fp))
        if not ok:
            failed[Path(fp).name] = reason

    if not failed:
        state["state"] = "CALCULATING"
        state["cleaned_files"] = state["uploaded_files"]
        sys_p = (
            "你是重譯文本分析系統的工作程序 Agent，"
            "請用繁體中文告知使用者所有檔案已確認為 ckip_segmentation_sumin 斷詞格式，"
            "即將開始計算指標，30 字以內。"
        )
        try:
            msg = call_worker_agent(sys_p, "所有檔案格式驗證通過。")
        except:
            msg = "[OK] 所有檔案均符合 ckip_segmentation_sumin 斷詞格式，即將開始計算文本指標……"
        return msg, state, False

    # 有檔案不符
    fail_list = "\n".join([f"- **{fn}**：{reason}" for fn, reason in failed.items()])
    sys_p = (
        "你是重譯文本分析系統的工作程序 Agent，請用繁體中文告知使用者："
        "部分檔案不符合 ckip_segmentation_sumin 斷詞工具的輸出格式，"
        "必須重新使用該工具斷詞後再上傳，語氣清晰，80 字以內。"
        "請附上工具連結：https://github.com/suhsiung/ckip_segmentation_sumin"
    )
    try:
        msg = call_worker_agent(sys_p, f"格式不符的檔案：\n{fail_list}")
        msg = msg + f"\n\n{fail_list}"
    except:
        msg = (
            f"[WARNING] 以下檔案不符合 ckip_segmentation_sumin 斷詞格式，請重新處理後再上傳：\n\n"
            f"{fail_list}\n\n"
            "請使用 [ckip_segmentation_sumin](https://github.com/suhsiung/ckip_segmentation_sumin) "
            "工具完成斷詞後，重新上傳檔案。"
        )
    state["uploaded_files"] = []
    state["state"] = "UPLOAD_WAIT"
    return msg, state, True


def run_computation_pipeline_streaming(state: dict):
    """
    執行完整計算流程（指標計算 → 統計分析 → LLM 報告），
    使用 generator 逐步 yield (progress_msg, is_final, state)。
    """
    output_dir  = Path(state["session_dir"]) / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    results_paths = []
    anal = None

    # ── 步驟 1：指標計算 ─────────────────────────────────────────────────────
    yield "【步驟 1/3】正在計算文本指標……", False, state
    try:
        idx_df, xlsx_path, csv_path = run_indicator_calculation(
            state["cleaned_files"], output_dir, _RESOURCES
        )
        results_paths.extend([str(xlsx_path), str(csv_path)])
        yield f"  [OK] 指標計算完成（{len(idx_df)} 個文本）", False, state
    except Exception as e:
        state["state"] = "COMPLETE"
        yield f"[ERROR] 指標計算失敗：{e}", True, state
        return

    # ── 步驟 2：統計分析 ─────────────────────────────────────────────────────
    yield "【步驟 2/3】正在執行統計分析……", False, state
    try:
        anal = run_statistical_analysis(csv_path, output_dir)
        for key in ["desc_path","diff_path","imp_path","imp_fig_path",
                    "box_path","dendro_path","elbow_path","cluster_fig_path",
                    "cluster_csv_path","report_path"]:
            p = anal.get(key)
            if p and Path(str(p)).exists():
                results_paths.append(str(p))
        state["analysis_result"] = {
            k: str(v) if isinstance(v, Path) else v
            for k, v in anal.items() if k != "df"
        }
        yield (
            f"  [OK] 統計分析完成（{anal['sig_count']}/{anal['total_indicators']} 個顯著指標，"
            f"隨機森林準確率 {anal['accuracy']:.4f}）"
        ), False, state
    except Exception as e:
        state["state"] = "COMPLETE"
        state["output_files"] = results_paths
        yield f"[ERROR] 統計分析失敗：{e}", True, state
        return

    # ── 步驟 3：LLM 研究報告 ────────────────────────────────────────────────
    yield "【步驟 3/3】主 Agent 正在撰寫研究報告（預計需要 30-90 秒）……", False, state
    try:
        rpt = run_report_writing(
            scenario_key  = state["scenario"],
            user_metadata = state.get("user_metadata", ""),
            user_focus    = state.get("user_focus", ""),
            csv_path      = csv_path,
            output_dir    = output_dir,
        )
        results_paths.extend([str(rpt["docx_path"]), str(rpt["txt_path"])])
        yield "  [OK] 研究報告撰寫完成", False, state
    except Exception as e:
        yield f"  [WARNING] LLM 報告撰寫失敗（{e}），其他分析結果仍可下載。", False, state

    # ── 完成 ─────────────────────────────────────────────────────────────────
    state["output_files"] = results_paths
    state["state"]        = "COMPLETE"

    sig_info = ""
    if anal:
        sig_info = (
            f"顯著指標：**{anal['sig_count']}/{anal['total_indicators']}** 個（p < {ALPHA}）。\n"
            f"隨機森林分類準確率：**{anal['accuracy']:.4f}**。\n"
        )
    summary = (
        f"**全部分析完成！**\n\n"
        f"共分析 **{len(state['cleaned_files'])}** 個文本。\n"
        f"{sig_info}\n"
        f"各項結果請點選上方各分頁查看。"
    )
    yield summary, True, state


# ============================================================================
# SECTION 10: Gradio Interface (Tabbed)
# ============================================================================

_EMPTY_HINT = (
    '<div style="text-align:center;padding:4rem 1rem 3rem;">'
    '<div style="width:64px;height:64px;margin:0 auto 1.25rem;border-radius:16px;'
    'background:rgba(14,165,233,.07);border:1px solid rgba(14,165,233,.15);'
    'display:flex;align-items:center;justify-content:center;">'
    '<svg xmlns="http://www.w3.org/2000/svg" width="28" height="28" fill="none" '
    'viewBox="0 0 24 24" stroke="rgba(14,165,233,.5)" stroke-width="1.5">'
    '<path stroke-linecap="round" stroke-linejoin="round" '
    'd="M3.75 12h16.5m-16.5 3.75h16.5M3.75 19.5h16.5M5.625 4.5h12.75a1.875 1.875 0 010 3.75H5.625a1.875 1.875 0 010-3.75z"/>'
    '</svg></div>'
    '<p style="font-size:1rem;font-weight:600;color:#94a3b8;margin:0 0 .4rem;">尚未產生資料</p>'
    '<p style="font-size:0.82rem;color:#475569;margin:0;">請先於「主頁」完成分析流程</p>'
    '</div>'
)

_CUSTOM_CSS = """
/* ═══════════════════════════════════════════════════════════════
   重譯文本分析系統 Agent — Refined Dark Academic Dashboard
   Palette: Sky-500 (#0ea5e9) + Indigo-400 (#818cf8) on Slate-950
   Typography: Inter (UI) + JetBrains Mono (data) + Crimson Pro (headings)
   WCAG AA compliant (4.5:1+)
   ═══════════════════════════════════════════════════════════════ */

@import url('https://fonts.googleapis.com/css2?family=Crimson+Pro:wght@400;600;700&display=swap');

/* ── 1. TOKENS & RESET ─────────────────────────────────────────── */
:root, .dark {
    --c-bg:           #020617;
    --c-surface:      #0d1526;
    --c-surface2:     #111827;
    --c-border:       rgba(14,165,233,.11);
    --c-border-hover: rgba(14,165,233,.30);
    --c-accent:       #0ea5e9;
    --c-accent2:      #818cf8;
    --c-accent3:      #c084fc;
    --c-text:         #e2e8f0;
    --c-text-muted:   #94a3b8;
    --c-text-subtle:  #475569;
    --c-success:      #34d399;
    --c-warn:         #fbbf24;

    --body-background-fill:      var(--c-bg) !important;
    --background-fill-primary:   var(--c-surface) !important;
    --background-fill-secondary: #090e1a !important;
    --block-background-fill:     var(--c-surface) !important;
    --block-border-color:        var(--c-border) !important;
    --border-color-primary:      var(--c-border) !important;
    --color-accent:              var(--c-accent) !important;
    --color-accent-soft:         rgba(14,165,233,.08) !important;
    --body-text-color:           var(--c-text) !important;
    --block-label-text-color:    var(--c-text-muted) !important;
    --block-title-text-color:    #cbd5e1 !important;
    --input-background-fill:     #1a2540 !important;
    --input-border-color:        rgba(14,165,233,.15) !important;
    --button-primary-background-fill: var(--c-accent) !important;
    --button-primary-text-color: #fff !important;
    --shadow-drop:               0 1px 3px rgba(0,0,0,.5) !important;
}
.gradio-container {
    max-width: 1680px !important;
    margin: 0 auto !important;
    padding: 0 1.5rem 1.5rem !important;
}
*, *::before, *::after { box-sizing: border-box; }

/* ── 2. HEADER（卡片化）─────────────────────────────────────────── */
.app-header {
    text-align: center;
    padding: 1.6rem 1.5rem 1.25rem;
    position: relative;
    margin: 1rem 0 1.25rem;
    border: 1px solid var(--c-border);
    border-radius: 16px;
    background: linear-gradient(180deg, rgba(14,165,233,.07) 0%, rgba(15,23,42,.45) 100%);
}
.app-header h1 {
    font-family: 'Crimson Pro', 'Noto Serif TC', Georgia, serif !important;
    font-size: 2.1rem !important;
    font-weight: 700 !important;
    background: linear-gradient(130deg, #38bdf8 0%, #818cf8 55%, #c084fc 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    margin: 0 0 .5rem !important;
    letter-spacing: -.01em;
    line-height: 1.2 !important;
}
.app-header .subtitle {
    font-size: 1.1rem !important;
    line-height: 1.5 !important;
    margin: 0 0 .25rem !important;
    background: linear-gradient(130deg, #38bdf8 0%, #818cf8 55%, #c084fc 100%) !important;
    -webkit-background-clip: text !important;
    -webkit-text-fill-color: transparent !important;
    background-clip: text !important;
    color: transparent !important;
}
/* 第三排：字型比第二排大、比主標題小 */
.app-header .subtitle-3 {
    font-size: 1.5rem !important;
    font-weight: 700 !important;
    line-height: 1.4 !important;
    margin: 0 !important;
    background: linear-gradient(130deg, #38bdf8 0%, #818cf8 55%, #c084fc 100%) !important;
    -webkit-background-clip: text !important;
    -webkit-text-fill-color: transparent !important;
    background-clip: text !important;
    color: transparent !important;
}
.app-header .subtitle strong {
    -webkit-text-fill-color: transparent !important;
    color: transparent !important;
    font-weight: 600;
}
.badge {
    display: inline-flex; align-items: center; gap: 4px;
    padding: 2px 9px;
    border-radius: 20px;
    font-size: 0.70rem;
    font-weight: 700;
    letter-spacing: 0.05em;
    margin: 0 3px;
    vertical-align: middle;
}
.badge-llm1 {
    background: rgba(14,165,233,.10);
    color: #38bdf8;
    border: 1px solid rgba(14,165,233,.22);
}
.badge-llm2 {
    background: rgba(129,140,248,.10);
    color: #a5b4fc;
    border: 1px solid rgba(129,140,248,.22);
}
/* step indicators */
.app-header .steps {
    display: flex; justify-content: center; flex-wrap: wrap;
    gap: 4px; margin-top: 1rem;
}
.step-pill {
    display: inline-flex; align-items: center; gap: 6px;
    padding: 4px 13px;
    border-radius: 20px;
    font-size: 0.72rem;
    font-weight: 600;
    background: rgba(14,165,233,.06);
    border: 1px solid rgba(56,189,248,.30);
    color: #cbd5e1;
    letter-spacing: .03em;
    transition: border-color .2s, background .2s;
}
.step-pill:hover {
    background: rgba(14,165,233,.12);
    border-color: rgba(56,189,248,.55);
}
.step-pill .step-num {
    width: 17px; height: 17px; border-radius: 50%;
    background: rgba(14,165,233,.22);
    color: #38bdf8;
    display: inline-flex; align-items: center; justify-content: center;
    font-size: 0.62rem; font-weight: 800;
}

/* ── 3. TABS ───────────────────────────────────────────────────── */
.tabs > .tab-nav {
    background: rgba(9,14,26,.92) !important;
    backdrop-filter: blur(16px) !important;
    -webkit-backdrop-filter: blur(16px) !important;
    border-bottom: 1px solid var(--c-border) !important;
    padding: 0 1rem !important;
    gap: 0 !important;
    position: sticky; top: 0; z-index: 50;
}
.tabs > .tab-nav > button {
    color: var(--c-text-subtle) !important;
    font-weight: 600 !important;
    font-size: 1.20rem !important;
    padding: 0.70rem 1.1rem !important;
    border: none !important;
    border-bottom: 2px solid transparent !important;
    border-radius: 0 !important;
    background: transparent !important;
    transition: color .18s ease, border-color .18s ease, background .18s ease !important;
    cursor: pointer !important;
    white-space: nowrap;
}
.tabs > .tab-nav > button:hover {
    color: #cbd5e1 !important;
    background: rgba(14,165,233,.05) !important;
}
.tabs > .tab-nav > button.selected {
    color: #38bdf8 !important;
    border-bottom-color: var(--c-accent) !important;
    background: rgba(14,165,233,.07) !important;
}
.tabitem { padding: 1.5rem 1.25rem !important; }
/* #7 語言指標分析等結果大標題置中 */
.result-md h1, .result-md h2, .result-md h3 { text-align: center !important; }
/* 巢狀子分頁（統計比較分析內）維持原字級 */
.tabitem .tabs > .tab-nav > button { font-size: 0.80rem !important; }

/* ── 4. CHATBOT ────────────────────────────────────────────────── */
.chatbot {
    border: 1px solid var(--c-border) !important;
    border-radius: 14px !important;
    background: #0f172a !important;
    backdrop-filter: blur(8px) !important;
}
/* 主對話視窗：統一深色 */
.main-chat {
    background: #0f172a !important;
    border: 1px solid #1e293b !important;
}
.main-chat .bubble-wrap,
.main-chat div[role="log"] {
    background: #0f172a !important;
}
/* #2 消除重複滾軸：只讓 bubble-wrap 捲動，外層不捲 */
.main-chat [class*="wrapper"],
.main-chat > div > div:not([role="log"]) {
    overflow: hidden !important;
}
.main-chat div[role="log"].bubble-wrap {
    overflow-y: auto !important;
    overflow-x: hidden !important;
}

/* 訊息從頂部開始排列，消除空白 */
div[role="log"].bubble-wrap,
div[role="log"][aria-label="chatbot conversation"],
.main-chat .bubble-wrap,
.bubble-wrap {
    flex-direction: column !important;
    justify-content: flex-start !important;
    align-content: flex-start !important;
    align-items: flex-start !important;
    min-height: 0 !important;
    padding: 0.5rem !important;
}
/* 清除 Gradio 在 bubble-wrap 內加的空白 spacer div */
.bubble-wrap > div:empty { display: none !important; }
.bubble-wrap > :first-child { margin-top: 0 !important; padding-top: 0 !important; }

/* ── 隱藏訊息計數氣泡、頁碼、複製按鈕浮層 ── */
.chatbot .count,
.chatbot .message-count,
.chatbot [class*="count"],
.chatbot .page,
.chatbot .page-count,
.chatbot .scroll-hide,
.chatbot .copy-text-button,
.chatbot .copy-btn,
.chatbot button[title*="copy" i],
.chatbot button[aria-label*="copy" i],
.chatbot [data-testid*="copy"],
.chatbot .bubble-wrap > div:last-child > span:only-child {
    display: none !important;
}
/* 徹底隱藏對話視窗內所有複製/工具浮層圖示 */
.main-chat .icon-button-wrapper,
.main-chat [class*="icon-button"],
.main-chat .message-buttons,
.main-chat .message-row button,
.main-chat button[title*="複製"],
.main-chat button[aria-label*="複製"] {
    display: none !important;
}

/* ── 通用訊息氣泡 ── */
.chatbot .message-row { margin-bottom: .5rem !important; }
.chatbot .message-row .message {
    border-radius: 10px !important;
    font-size: 0.80rem !important;
    line-height: 1.7 !important;
    padding: .6rem .95rem !important;
    max-width: 90% !important;
    transition: box-shadow .2s !important;
}

/* 使用者訊息（所有 chatbot：綠色背景、緊湊間距） */
.chatbot .message-row.user .message {
    background: rgba(34,197,94,.18) !important;
    border: 1px solid rgba(34,197,94,.42) !important;
    color: #bbf7d0 !important;
    margin-left: auto !important;
    font-weight: 500 !important;
    box-shadow: none !important;
    padding: .4rem .8rem !important;
    width: fit-content !important;
    max-width: 80% !important;
}
.chatbot .message-row.user { margin-bottom: .35rem !important; }

/* AI 訊息（主對話視窗） */
.main-chat .message-row.bot .message {
    background: #1e293b !important;
    border: 1px solid #334155 !important;
    border-left: 3px solid #38bdf8 !important;
    color: #e2e8f0 !important;
    font-size: 0.80rem !important;
    line-height: 1.7 !important;
    box-shadow: 0 1px 3px rgba(0,0,0,.3) !important;
}
.main-chat .message-row.bot .message strong,
.main-chat .message-row.bot .message b { color: #38bdf8 !important; font-weight: 700 !important; }
.main-chat .message-row.bot .message em,
.main-chat .message-row.bot .message i  { color: #a5b4fc !important; }
.main-chat .message-row.bot .message code {
    background: rgba(14,165,233,.12) !important;
    color: #fbbf24 !important;
    padding: 1px 6px !important;
    border-radius: 4px !important;
    font-family: 'JetBrains Mono', monospace !important;
    font-size: 0.82em !important;
}
.main-chat .message-row.bot .message a {
    color: #38bdf8 !important;
    text-decoration: underline !important;
    text-underline-offset: 2px !important;
}

/* 執行視窗（簡易訊息顯示欄，非對話氣泡） */
.exec-log {
    background: rgba(8,12,26,.6) !important;
    border: 1px solid var(--c-border) !important;
    border-radius: 10px !important;
    padding: .55rem .75rem !important;
    max-height: 300px !important;
    overflow-y: auto !important;
    margin-top: .6rem !important;
}
.exec-log p {
    font-size: 0.99rem !important;
    line-height: 1.5 !important;
    margin: 0 0 .3rem !important;
    padding: 0 !important;
    color: #94a3b8 !important;
    font-family: 'JetBrains Mono', 'Noto Sans TC', monospace !important;
    white-space: pre-wrap !important;
    word-break: break-word !important;
}
.exec-log p:first-child { color: #cbd5e1 !important; }
.exec-log code {
    background: none !important;
    color: #38bdf8 !important;
    padding: 0 .15rem 0 0 !important;
    font-size: 0.99rem !important;
}
.exec-log::-webkit-scrollbar { width: 6px; }
.exec-log::-webkit-scrollbar-thumb { background: rgba(56,189,248,.2); border-radius: 3px; }

/* ── 5. TODO PANEL ────────────────────────────────────────────── */
.todo-panel {
    background: rgba(8,12,26,.75) !important;
    border: 1px solid var(--c-border) !important;
    border-radius: 12px !important;
    padding: .9rem 1rem !important;
    margin-bottom: .75rem !important;
}
.todo-panel p, .todo-panel div { font-size: 1.20rem !important; line-height: 1.8 !important; }
.todo-panel strong { color: #38bdf8 !important; font-weight: 700 !important; }
.todo-panel code {
    background: rgba(14,165,233,.08) !important;
    color: #34d399 !important;
    border-radius: 4px !important;
    font-size: 1.08rem !important;
    letter-spacing: .03em !important;
    padding: 1px 5px !important;
}
/* ── 5b. COLUMN LAYOUT ─────────────────────────────────────────── */
.chat-col { display: flex !important; flex-direction: column !important; }
.side-col { display: flex !important; flex-direction: column !important; }
/* 執行進度（to-do-list）：填滿右欄整個高度 */
.side-col .todo-panel {
    flex: 1 1 auto !important;
    min-height: 740px !important;
}

/* ── 5c. INPUT ROW ─────────────────────────────────────────────── */
.input-row {
    margin-top: .6rem !important;
    gap: .6rem !important;
    display: flex !important;
    flex-wrap: nowrap !important;
    align-items: stretch !important;
    width: 100% !important;
}
/* 輸入框：撐滿剩餘空間 */
.input-row > *:first-child {
    flex: 1 1 auto !important;
    min-width: 0 !important;
}
/* 送出按鈕：固定寬度貼右緣 */
.input-row > *:last-child {
    flex: 0 0 140px !important;
    min-width: 140px !important;
}
.input-row textarea,
.input-row input[type="text"] {
    background: #1e293b !important;
    border: 1px solid #334155 !important;
    color: #e2e8f0 !important;
    border-radius: 10px !important;
    font-size: 0.88rem !important;
    padding: .7rem 1rem !important;
    transition: border-color .2s !important;
}
.input-row textarea:focus,
.input-row input[type="text"]:focus {
    border-color: #38bdf8 !important;
    outline: none !important;
    box-shadow: 0 0 0 2px rgba(56,189,248,.15) !important;
}
/* 送出按鈕樣式 */
.input-row button {
    height: 100% !important;
    min-height: 46px !important;
    border-radius: 10px !important;
    font-weight: 700 !important;
    font-size: 0.95rem !important;
}
/* ── 6. UPLOAD AREA（醒目卡片）─────────────────────────────────── */
.upload-zone {
    border: 1px solid rgba(56,189,248,.40) !important;
    border-radius: 16px !important;
    background: linear-gradient(180deg, rgba(14,165,233,.06) 0%, #0f172a 100%) !important;
    padding: 1.4rem 1.5rem !important;
    margin: 1rem auto !important;
    max-width: 720px !important;
    box-shadow: 0 16px 50px rgba(0,0,0,.45) !important;
    animation: popupIn .25s cubic-bezier(.2,.8,.3,1) !important;
}
@keyframes popupIn {
    from { opacity: 0; transform: translateY(-8px) scale(.985); }
    to   { opacity: 1; transform: translateY(0) scale(1); }
}
.upload-zone::before {
    content: "📂 上傳已斷詞文本";
    display: block;
    font-size: 1.05rem;
    font-weight: 700;
    color: #38bdf8;
    margin-bottom: .9rem;
    letter-spacing: .02em;
}
.upload-zone:hover, .upload-zone:focus-within {
    border-color: rgba(56,189,248,.6) !important;
}
.upload-actions {
    margin-top: 1rem !important;
    gap: .6rem !important;
}
.upload-actions button { border-radius: 10px !important; font-weight: 700 !important; }
/* 語料 Metadata 說明表單卡片 */
.metadata-zone {
    border: 1px solid rgba(56,189,248,.40) !important;
    border-radius: 16px !important;
    background: linear-gradient(180deg, rgba(14,165,233,.06) 0%, #0f172a 100%) !important;
    padding: 1.4rem 1.6rem !important;
    margin: 1rem 0 !important;
    box-shadow: 0 16px 50px rgba(0,0,0,.45) !important;
    animation: popupIn .25s cubic-bezier(.2,.8,.3,1) !important;
}
.metadata-zone .gr-box, .metadata-zone .form { background: transparent !important; }
.upload-area, [data-testid="droparea"] {
    border: 2px dashed rgba(14,165,233,.22) !important;
    border-radius: 12px !important;
    background: rgba(14,165,233,.025) !important;
    transition: border-color .2s ease, background .2s ease !important;
}
.upload-area:hover, [data-testid="droparea"]:hover {
    border-color: rgba(14,165,233,.45) !important;
    background: rgba(14,165,233,.06) !important;
}
.file-preview {
    border: 1px solid var(--c-border) !important;
    border-radius: 10px !important;
    background: var(--c-surface) !important;
}

/* ── 6. BUTTONS ────────────────────────────────────────────────── */
button.primary, button[data-testid="submit"] {
    background: linear-gradient(135deg, #0284c7 0%, #4f46e5 100%) !important;
    border: none !important;
    border-radius: 9px !important;
    font-weight: 700 !important;
    font-size: 0.84rem !important;
    color: #fff !important;
    box-shadow: 0 2px 12px rgba(14,165,233,.22) !important;
    transition: transform .18s ease, box-shadow .18s ease, opacity .18s !important;
    cursor: pointer !important;
    letter-spacing: .02em;
}
button.primary:hover { transform: translateY(-2px) !important; box-shadow: 0 6px 22px rgba(14,165,233,.35) !important; }
button.primary:active { transform: translateY(0) !important; opacity: .92 !important; }
button.primary:disabled { opacity: .5 !important; cursor: not-allowed !important; transform: none !important; }
button.secondary {
    background: rgba(14,165,233,.07) !important;
    border: 1px solid rgba(14,165,233,.22) !important;
    color: #38bdf8 !important;
    border-radius: 9px !important;
    font-weight: 600 !important;
    font-size: 0.84rem !important;
    cursor: pointer !important;
    transition: background .18s ease, border-color .18s ease !important;
}
button.secondary:hover {
    background: rgba(14,165,233,.14) !important;
    border-color: rgba(14,165,233,.40) !important;
}

/* ── 7. INPUTS ─────────────────────────────────────────────────── */
textarea, input[type="text"] {
    background: #1a2540 !important;
    border: 1px solid rgba(14,165,233,.14) !important;
    border-radius: 9px !important;
    color: var(--c-text) !important;
    font-size: 0.875rem !important;
    line-height: 1.6 !important;
    transition: border-color .18s ease, box-shadow .18s ease !important;
}
textarea:focus, input[type="text"]:focus {
    border-color: var(--c-accent) !important;
    box-shadow: 0 0 0 3px rgba(14,165,233,.14) !important;
    outline: none !important;
}

/* ── 8. DATAFRAME ──────────────────────────────────────────────── */
.table-wrap {
    border: 1px solid var(--c-border) !important;
    border-radius: 12px !important;
    overflow: hidden !important;
}
table thead tr { background: #111827 !important; }
table thead th {
    background: transparent !important;
    color: var(--c-text-muted) !important;
    font-weight: 700 !important;
    font-size: 0.75rem !important;
    text-transform: uppercase !important;
    letter-spacing: 0.05em !important;
    border-bottom: 1px solid rgba(14,165,233,.12) !important;
    padding: 10px 12px !important;
    white-space: nowrap;
}
table tbody td {
    background: var(--c-surface) !important;
    color: #cbd5e1 !important;
    font-size: 0.80rem !important;
    border-color: rgba(14,165,233,.04) !important;
    font-family: 'JetBrains Mono', monospace !important;
    padding: 7px 12px !important;
}
table tbody tr:nth-child(even) td { background: rgba(14,165,233,.025) !important; }
table tbody tr:hover td { background: rgba(14,165,233,.06) !important; }

/* ── 9. IMAGE FRAMES ───────────────────────────────────────────── */
.image-container, .image-frame {
    border: 1px solid var(--c-border) !important;
    border-radius: 12px !important;
    overflow: hidden !important;
    background: var(--c-surface) !important;
}
.image-container img { display: block; width: 100%; height: auto; }

/* ── 10. RESULT SECTION CARDS ──────────────────────────────────── */
.result-card {
    background: var(--c-surface);
    border: 1px solid var(--c-border);
    border-radius: 14px;
    padding: 1.25rem 1.5rem;
    margin-bottom: 1.25rem;
}
.result-card-title {
    font-size: 0.72rem;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: .08em;
    color: var(--c-text-subtle);
    margin: 0 0 .75rem;
    padding-bottom: .5rem;
    border-bottom: 1px solid var(--c-border);
    display: flex; align-items: center; gap: 6px;
}
.result-card-title .dot {
    width: 6px; height: 6px; border-radius: 50%;
    background: var(--c-accent);
    display: inline-block;
}
.stat-pill {
    display: inline-flex; align-items: center; gap: 5px;
    background: rgba(14,165,233,.08);
    border: 1px solid rgba(14,165,233,.15);
    border-radius: 6px;
    padding: 3px 10px;
    font-size: 0.78rem;
    color: #38bdf8;
    font-weight: 600;
    margin: 2px;
}

/* ── 11. MARKDOWN (result tabs) ────────────────────────────────── */
.prose h3, .md h3, .svelte-1ipelgc h3 {
    font-family: 'Crimson Pro', Georgia, serif !important;
    font-size: 1.25rem !important;
    font-weight: 600 !important;
    color: #e2e8f0 !important;
    border-bottom: 1px solid var(--c-border);
    padding-bottom: 0.45rem;
    margin: 1.4rem 0 .8rem !important;
}
.prose h4, .md h4 {
    font-size: 0.85rem !important;
    font-weight: 700 !important;
    color: var(--c-text-muted) !important;
    text-transform: uppercase;
    letter-spacing: .06em;
    margin: 1rem 0 .4rem !important;
}
.prose strong, .md strong { color: #38bdf8 !important; }
.prose em, .md em { color: #a5b4fc !important; }
.prose code, .md code {
    background: #1a2540 !important;
    color: #fbbf24 !important;
    padding: 2px 6px;
    border-radius: 4px;
    font-family: 'JetBrains Mono', monospace !important;
    font-size: 0.82em;
}
.prose ul li, .md ul li { color: #cbd5e1; line-height: 1.75; }
.prose p { line-height: 1.75; color: #cbd5e1; }

/* ── 12. LABELS ────────────────────────────────────────────────── */
.block > .wrap > label > span,
.block > .label-wrap > span,
label > span {
    color: var(--c-text-subtle) !important;
    font-weight: 700 !important;
    font-size: 0.70rem !important;
    text-transform: uppercase !important;
    letter-spacing: 0.08em !important;
}

/* ── 13. DOWNLOAD SECTION ──────────────────────────────────────── */
.download-bar {
    display: flex; align-items: center; gap: 10px;
    background: linear-gradient(135deg,rgba(14,165,233,.06),rgba(99,102,241,.06));
    border: 1px solid rgba(14,165,233,.15);
    border-radius: 12px;
    padding: .8rem 1.1rem;
    margin-top: 1rem;
}
.download-bar svg { flex-shrink: 0; }

/* ── 14. SCROLLBAR ─────────────────────────────────────────────── */
::-webkit-scrollbar { width: 5px; height: 5px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: #1e3a5f; border-radius: 3px; }
::-webkit-scrollbar-thumb:hover { background: #334d6e; }

/* ── 15. FOCUS & ACCESSIBILITY ─────────────────────────────────── */
:focus-visible {
    outline: 2px solid var(--c-accent) !important;
    outline-offset: 2px !important;
    border-radius: 4px;
}
a { color: #38bdf8; transition: color .15s; }
a:hover { color: #7dd3fc; }

/* ── 16. REDUCED MOTION ────────────────────────────────────────── */
@media (prefers-reduced-motion: reduce) {
    *, *::before, *::after {
        animation-duration: .01ms !important;
        transition-duration: .01ms !important;
    }
}

/* ── 16b. LEGEND BAR ───────────────────────────────────────────── */
.legend-bar {
    display: flex; flex-wrap: wrap;
    justify-content: center; align-items: center;
    gap: 1.75rem;
    margin-top: 1.25rem;
    padding: .85rem 1.25rem;
    border: 1px solid var(--c-border);
    border-radius: 12px;
    background: rgba(15,23,42,.5);
}
.legend-item {
    display: inline-flex; align-items: center; gap: .5rem;
    font-size: 0.78rem;
    color: var(--c-text-subtle);
}
.legend-item .dot {
    width: 9px; height: 9px; border-radius: 50%;
    display: inline-block; flex-shrink: 0;
}
.dot-blue  { background: #38bdf8; }
.dot-amber { background: #fbbf24; }
.dot-green { background: #34d399; }

/* ── 17. RESPONSIVE ────────────────────────────────────────────── */
@media (max-width: 768px) {
    .app-header h1 { font-size: 1.6rem !important; }
    .tabitem { padding: 1rem .75rem !important; }
    .app-header .steps { gap: 3px; }
    .step-pill { font-size: 0.62rem; padding: 2px 7px; }
}
"""


def create_app():
    with gr.Blocks(
        title="重譯文本分析系統Agent",
        theme=gr.themes.Base(
            primary_hue=gr.themes.colors.sky,
            secondary_hue=gr.themes.colors.indigo,
            neutral_hue=gr.themes.colors.slate,
            font=[gr.themes.GoogleFont("Inter"), "Noto Sans TC", "system-ui", "sans-serif"],
            font_mono=[gr.themes.GoogleFont("JetBrains Mono"), "monospace"],
        ),
        css=_CUSTOM_CSS,
        js="""() => {
    document.body.classList.add('dark');
    document.title = '重譯文本分析系統Agent';

    function fixBubbleWrap() {
        /* 直接鎖定 .bubble-wrap，強制 column + flex-start */
        document.querySelectorAll('.main-chat .bubble-wrap, .main-chat div[role="log"]').forEach(function(el) {
            el.style.setProperty('flex-direction', 'column', 'important');
            el.style.setProperty('justify-content', 'flex-start', 'important');
            el.style.setProperty('align-items', 'flex-start', 'important');
            el.style.setProperty('align-content', 'flex-start', 'important');
            el.scrollTop = 0;
        });
    }

    /* 多次觸發確保 Svelte hydration 後生效 */
    [300, 800, 1500, 3000].forEach(function(t) { setTimeout(fixBubbleWrap, t); });

    /* 持續監聽 DOM 變化（每次訊息更新後重設） */
    var mo = new MutationObserver(function(mutations) {
        var needsFix = mutations.some(function(m) {
            return m.target.closest && m.target.closest('.main-chat');
        });
        if (needsFix) { fixBubbleWrap(); }
    });
    setTimeout(function() {
        var root = document.querySelector('.main-chat');
        if (root) mo.observe(root, { childList: true, subtree: true, attributes: true });
    }, 500);

    /* metadata 表單：出現時自動捲入視野；消失時捲回頂部 */
    var metaWasVisible = false;
    function watchMetadata() {
        var zone = document.getElementById('metadata_zone');
        var visible = !!(zone && zone.offsetParent !== null);
        if (visible && !metaWasVisible) {
            zone.scrollIntoView({ behavior: 'smooth', block: 'center' });
        } else if (!visible && metaWasVisible) {
            window.scrollTo({ top: 0, behavior: 'smooth' });
        }
        metaWasVisible = visible;
    }
    setInterval(watchMetadata, 400);
}""",
    ) as demo:

        # ── Header ──────────────────────────────────────────────────────
        gr.HTML(
            '<div class="app-header">'
            '  <h1>譯彩紛呈：重譯文本分析系統</h1>'
            '  <p class="subtitle">TransPrism: An Analytical Framework for Retranslation</p>'
            '  <p class="subtitle-3">II 分析引擎 SpectraEngine</p>'
            '  <div class="steps">'
            '    <span class="step-pill"><span class="step-num">1</span>情境選擇</span>'
            '    <span class="step-pill"><span class="step-num">2</span>上傳引導</span>'
            '    <span class="step-pill"><span class="step-num">3</span>文本上傳</span>'
            '    <span class="step-pill"><span class="step-num">4</span>格式驗證</span>'
            '    <span class="step-pill"><span class="step-num">5</span>指標計算</span>'
            '    <span class="step-pill"><span class="step-num">6</span>統計分析</span>'
            '    <span class="step-pill"><span class="step-num">7</span>報告撰寫</span>'
            '    <span class="step-pill"><span class="step-num">8</span>下載成果</span>'
            '  </div>'
            '</div>'
        )

        session_state = gr.State(None)

        with gr.Tabs():
            # ══════════════════════════════════════════════════════════════
            # Tab 0: 主頁
            # ══════════════════════════════════════════════════════════════
            with gr.Tab("🏠 主頁"):
                with gr.Row(equal_height=False):
                    # ── 左欄：對話視窗 → 輸入列 → 執行視窗（堆疊）────────
                    with gr.Column(scale=7, min_width=460, elem_classes=["chat-col"]):
                        chatbot = gr.Chatbot(
                            label="對話視窗",
                            height=400,
                            elem_classes=["main-chat"],
                        )
                        with gr.Row(elem_classes=["input-row"]):
                            msg_input = gr.Textbox(
                                placeholder="輸入訊息並按 Enter 或點擊「送出」……",
                                show_label=False, scale=5,
                            )
                            send_btn = gr.Button("送出", scale=1, variant="primary")
                        exec_log = gr.Markdown(
                            value="",
                            label="執行視窗",
                            elem_classes=["exec-log"],
                        )
                    # ── 右欄：執行進度（to-do-list，整欄滿高）────────────
                    with gr.Column(scale=3, min_width=280, elem_classes=["side-col"]):
                        todo_md = gr.Markdown(
                            value="",
                            elem_classes=["todo-panel"],
                            label="執行進度",
                            visible=False,
                        )

                with gr.Column(visible=False, elem_classes=["upload-zone"]) as upload_row:
                    file_input = gr.File(
                        label="上傳已斷詞文本（.txt，可多選）",
                        file_types=[".txt"], file_count="multiple",
                    )
                    with gr.Row(elem_classes=["upload-actions"]):
                        cancel_btn = gr.Button("取消", scale=1, variant="secondary")
                        upload_btn = gr.Button("確認上傳", scale=2, variant="primary")

                # ── 語料 Metadata 說明表單（上傳後、分析前）──────────────────
                with gr.Column(visible=False, elem_id="metadata_zone", elem_classes=["metadata-zone"]) as metadata_row:
                    gr.Markdown("### 📑 語料 Metadata 說明\n請於正式分析前說明語料來源、範圍與比較結構，將寫入研究報告。")

                    gr.Markdown("**一、原文資訊**")
                    with gr.Row():
                        meta_title   = gr.Textbox(label="原文標題", scale=2)
                        meta_author  = gr.Textbox(label="作者", scale=2)
                        meta_lang    = gr.Textbox(label="語言", scale=1)
                    with gr.Row():
                        meta_pubyear = gr.Textbox(label="出版年", scale=1)
                        meta_version = gr.Textbox(label="版本", scale=2)

                    # 比較結構（僅情境 3、4 需要；情境 1、2 固定二群）
                    with gr.Group(visible=False) as meta_compare_group:
                        gr.Markdown("**二、比較結構**")
                        with gr.Row():
                            meta_ngroups = gr.Radio(
                                ["二群", "三群"], label="群組數量", value="二群", scale=1,
                            )
                            meta_method  = gr.Radio(
                                ["群組間比較", "個體間比較"], label="比較方式",
                                value="群組間比較", scale=1,
                            )

                    gr.Markdown("**三、群組層資訊**（共通屬性）")
                    meta_group_df = gr.Dataframe(
                        headers=["群組標籤", "翻譯類型", "翻譯方向", "譯者背景", "目標讀者"],
                        datatype=["str"] * 5,
                        col_count=(5, "fixed"),
                        row_count=(1, "dynamic"),
                        interactive=True,
                        label="群組層",
                    )

                    gr.Markdown("**四、個體層資訊**（每個譯本／檔案）")
                    meta_indiv_df = gr.Dataframe(
                        headers=["檔案", "譯者／系統名稱", "出版年", "譯者母語", "後編輯程度", "重譯序位"],
                        datatype=["str"] * 6,
                        col_count=(6, "fixed"),
                        row_count=(1, "dynamic"),
                        interactive=True,
                        label="個體層",
                    )

                    with gr.Row(elem_classes=["upload-actions"]):
                        meta_cancel_btn  = gr.Button("略過", scale=1, variant="secondary")
                        meta_confirm_btn = gr.Button("確認並開始分析", scale=2, variant="primary")

            # ══════════════════════════════════════════════════════════════
            # 結果分頁：初始隱藏，分析完成後顯示
            # ══════════════════════════════════════════════════════════════
            with gr.Tab("📊 語言指標分析", visible=False) as tab2:
                tab_idx_md   = gr.Markdown(_EMPTY_HINT, elem_classes=["result-md"])
                tab_idx_df   = gr.Dataframe(
                    label="指標計算結果",
                    interactive=False, visible=False,
                )
                tab_idx_file = gr.File(
                    label="下載檔案",
                    file_count="multiple", interactive=False, visible=False,
                )

            with gr.Tab("📈 統計比較分析", visible=False) as tab3:
                with gr.Tabs():
                    # 子分頁 1：敘述統計
                    with gr.Tab("1. 敘述統計"):
                        tab_desc_md   = gr.Markdown(_EMPTY_HINT)
                        tab_desc_df   = gr.Dataframe(
                            label="敘述性統計",
                            interactive=False, visible=False,
                        )
                    # 子分頁 2：差異性分析
                    with gr.Tab("2. 差異性分析"):
                        tab_diff_head = gr.Markdown(_EMPTY_HINT)
                        tab_diff_df   = gr.Dataframe(
                            label="差異性分析",
                            interactive=False, visible=False,
                        )
                        tab_box_img   = gr.Image(label="顯著指標箱形圖", visible=False)
                    # 子分頁 3：監督式學習
                    with gr.Tab("3. 監督式學習"):
                        tab_rf_md  = gr.Markdown(_EMPTY_HINT)
                        with gr.Row():
                            with gr.Column(scale=1):
                                tab_imp_df = gr.Dataframe(
                                    label="特徵重要性 Top 20",
                                    interactive=False, visible=False,
                                )
                            with gr.Column(scale=2):
                                tab_imp_img = gr.Image(label="特徵重要性圖", visible=False)
                    # 子分頁 4：非監督式學習
                    with gr.Tab("4. 非監督式學習"):
                        tab_clust_md = gr.Markdown(_EMPTY_HINT)
                        with gr.Row():
                            with gr.Column(scale=1):
                                tab_dendro_img = gr.Image(label="層次分群樹狀圖", visible=False)
                            with gr.Column(scale=1):
                                tab_elbow_img  = gr.Image(label="K-means Elbow Curve", visible=False)
                        tab_pca_img    = gr.Image(label="分群視覺化（PCA 2D）", visible=False)
                        tab_cluster_df = gr.Dataframe(
                            label="分群結果",
                            interactive=False, visible=False,
                        )
                    # 子分頁 5：其他
                    with gr.Tab("5. 下載"):
                        tab_stats_file = gr.File(
                            label="下載統計檔案",
                            file_count="multiple", interactive=False, visible=False,
                        )

            with gr.Tab("📝 LLM輔助分析", visible=False) as tab6:
                tab_rpt_md   = gr.Markdown(_EMPTY_HINT)
                tab_rpt_file = gr.File(
                    label="下載報告",
                    file_count="multiple", interactive=False, visible=False,
                )

        # ── 底部圖例列 ────────────────────────────────────────────────────
        gr.HTML(
            '<div class="legend-bar">'
            '  <span class="legend-item"><span class="dot dot-blue"></span>對話區：統一深色主題，AI 訊息左邊藍線區分</span>'
            '  <span class="legend-item"><span class="dot dot-amber"></span>進度面板：情境選擇後才顯示</span>'
            '  <span class="legend-item"><span class="dot dot-green"></span>執行紀錄：最新在上</span>'
            '</div>'
        )

        # ── 分頁 Tab 元件（控制顯示/隱藏）─────────────────────────────────
        _RESULT_TABS = [tab2, tab3, tab6]

        # ── 所有分頁內容輸出元件 + 3 個 Tab 可見性 ────────────────────────
        _TAB_OUTPUTS = [
            tab_idx_md, tab_idx_df, tab_idx_file,               # 0-2  語言指標分析
            tab_desc_md, tab_desc_df,                            # 3-4  敘述統計
            tab_diff_head, tab_diff_df, tab_box_img,             # 5-7  差異性分析
            tab_rf_md, tab_imp_df, tab_imp_img,                  # 8-10 監督式學習
            tab_clust_md, tab_dendro_img, tab_elbow_img,         # 11-13 非監督式學習
            tab_pca_img, tab_cluster_df,                         # 14-15
            tab_stats_file,                                      # 16  其他
            tab_rpt_md, tab_rpt_file,                            # 17-18 研究報告
            tab2, tab3, tab6,                                    # 19-21 (visibility)
        ]
        _N_TAB = len(_TAB_OUTPUTS)

        # ── helper ──────────────────────────────────────────────────────────
        def _log_append(log_history, msg):
            # 最新訊息放最前面（置頂顯示），以 markdown 字串呈現
            prev = log_history or ""
            line = f"`›` {msg}"
            return f"{line}\n\n{prev}" if prev else line

        def _populate_tabs(state):
            """streaming 結束後呼叫，讀取 output 檔案並填入各分頁。"""
            R = [gr.update()] * _N_TAB          # 預設全部不動
            if state is None:
                return R

            output_dir = Path(state["session_dir"]) / "output"
            uploaded   = state.get("uploaded_files", [])

            # ── 需要 output 目錄存在 ──────────────────────────────────────
            if not output_dir.exists():
                return R

            # ── 語言指標分析（index 0-2）──────────────────────────────────
            csv_path  = output_dir / "textIndex_calculated.csv"
            xlsx_list = sorted(output_dir.glob("textIndex_calculated_1_*.xlsx"))
            if csv_path.exists():
                df = pd.read_csv(str(csv_path))
                # 指標數 = 總欄位數扣除中介欄（Text/Type1/Type2）
                meta_cols = [c for c in ["Text", "Type1", "Type2"] if c in df.columns]
                n_indicators = len(df.columns) - len(meta_cols)
                R[0] = f"### 指標計算結果\n共 **{len(df)}** 個文本，**{n_indicators}** 個語言指標。"
                R[1] = gr.update(value=df, visible=True)
                dl = [str(csv_path)]
                if xlsx_list:
                    dl.append(str(xlsx_list[-1]))
                R[2] = gr.update(value=dl, visible=True)

            # ── 統計比較分析：敘述統計 + 差異分析（index 3-7）─────────────
            desc_path = output_dir / "01_descriptive_stats.csv"
            diff_path = output_dir / "02_difference_analysis.csv"
            box_path  = output_dir / "02_significant_boxplots.png"

            if desc_path.exists():
                desc = pd.read_csv(str(desc_path))
                R[3] = f"### 敘述統計\n共 **{len(desc)}** 個指標的敘述統計。"
                R[4] = gr.update(value=desc, visible=True)

            if diff_path.exists():
                diff = pd.read_csv(str(diff_path))
                sig  = diff[diff.get('顯著', pd.Series()) != 'ns'] if '顯著' in diff.columns else pd.DataFrame()
                R[5] = gr.update(
                    value=f"### 差異性分析\n共 **{len(diff)}** 個指標，其中 **{len(sig)}** 個達顯著差異（p < {ALPHA}）。",
                    visible=True,
                )
                R[6] = gr.update(value=diff, visible=True)

            if box_path.exists():
                R[7] = gr.update(value=str(box_path), visible=True)

            # ── 統計比較分析：監督式學習（index 8-10）─────────────────────
            imp_path = output_dir / "03_feature_importance.csv"
            imp_fig  = output_dir / "03_feature_importance.png"
            anal     = state.get("analysis_result") or {}

            if anal.get("accuracy") is not None:
                acc = anal["accuracy"]
                acc_str = f"{acc:.4f}" if isinstance(acc, float) else str(acc)
                R[8] = f"### 隨機森林分類結果\n模型準確率（Accuracy）：**{acc_str}**"

            if imp_path.exists():
                imp = pd.read_csv(str(imp_path)).head(20)
                R[9] = gr.update(value=imp, visible=True)
            if imp_fig.exists():
                R[10] = gr.update(value=str(imp_fig), visible=True)

            # ── 統計比較分析：非監督式學習（index 11-15）──────────────────
            dendro_path  = output_dir / "04a_dendrogram.png"
            elbow_path   = output_dir / "04b_kmeans_elbow.png"
            pca_fig_path = output_dir / "04c_clustering_pca2d.png"
            cluster_path = output_dir / "05_cluster_results.csv"

            pca_ev = anal.get("pca_ev", [])
            ev_str = ""
            if len(pca_ev) >= 2:
                ev_str = f"\n\nPCA 解釋變異量：PC1 = {pca_ev[0]:.1%}，PC2 = {pca_ev[1]:.1%}"
            if any(p.exists() for p in [dendro_path, elbow_path, pca_fig_path]):
                R[11] = f"### 集群分析結果{ev_str}"

            if dendro_path.exists():
                R[12] = gr.update(value=str(dendro_path), visible=True)
            if elbow_path.exists():
                R[13] = gr.update(value=str(elbow_path), visible=True)
            if pca_fig_path.exists():
                R[14] = gr.update(value=str(pca_fig_path), visible=True)
            if cluster_path.exists():
                cdf = pd.read_csv(str(cluster_path))
                R[15] = gr.update(value=cdf, visible=True)

            # ── 統計比較分析：其他（下載統計檔案，index 16）───────────────
            stat_files = [str(p) for p in [desc_path, diff_path] if p.exists()]
            if stat_files:
                R[16] = gr.update(value=stat_files, visible=True)

            # ── 研究報告（index 17-18）────────────────────────────────────
            txt_files  = sorted(output_dir.glob("LLM_Analysis-*.txt"), reverse=True)
            docx_files = sorted(output_dir.glob("LLM_Analysis-*.docx"), reverse=True)
            final_rpts = sorted(output_dir.glob("Final_Report-*.docx"), reverse=True)

            if txt_files:
                try:
                    with open(str(txt_files[0]), 'r', encoding='utf-8') as f:
                        rpt_text = f.read()
                    R[17] = f"### LLM 深度研究報告\n\n{rpt_text}"
                except Exception:
                    R[17] = "### LLM 深度研究報告\n\n（讀取報告時發生錯誤）"

            dl_rpt = []
            if final_rpts:
                dl_rpt.append(str(final_rpts[0]))
            if docx_files:
                dl_rpt.append(str(docx_files[0]))
            if txt_files:
                dl_rpt.append(str(txt_files[0]))
            if dl_rpt:
                R[18] = gr.update(value=dl_rpt, visible=True)

            # ── 顯示各分頁 Tab（index 19-21）──────────────────────────────
            if csv_path.exists():
                R[19] = gr.update(visible=True)      # 語言指標分析
            if desc_path.exists() or imp_path.exists() \
               or any(p.exists() for p in [dendro_path, elbow_path, pca_fig_path]):
                R[20] = gr.update(visible=True)      # 統計比較分析
            if txt_files or docx_files:
                R[21] = gr.update(visible=True)      # 研究報告

            return R

        # ── 啟動：先顯示「啟動中」，再載入 LLM 歡迎訊息 ───────────────────
        def on_load_init():
            """第一步：立即顯示啟動中提示。"""
            sid   = str(uuid.uuid4())[:8]
            state = make_initial_state(sid)
            return (
                [_assistant_msg("系統啟動中...")],
                "`›` 系統啟動中…",
                state,
                gr.update(visible=False),
                gr.update(value=""),
            )

        def on_load_ready(chat_history, log_history, state):
            """第二步：顯示固定歡迎訊息。"""
            intro = (
                "歡迎使用 TransPrism II 分析引擎 SpectraEngine。\n\n"
                "請依照您的研究設計選擇分析情境：\n\n"
                "1 人工翻譯 vs 機器翻譯\n"
                "2 全譯 vs 編譯／節譯／改寫\n"
                "3 同期不同譯者譯本比較\n"
                "4 不同時期重譯本的歷時比較\n\n"
                "輸入 1、2、3 或 4 。"
            )
            return (
                [_assistant_msg(intro)],
                "`›` 系統已就緒，等待使用者操作。",
            )

        demo.load(
            fn=on_load_init,
            outputs=[chatbot, exec_log, session_state, upload_row, todo_md],
        ).then(
            fn=on_load_ready,
            inputs=[chatbot, exec_log, session_state],
            outputs=[chatbot, exec_log],
        )

        # ── todo 即時更新 helper ─────────────────────────────────────────────
        def _todo_update(state):
            """根據 state 的 done_steps / current_step 回傳 todo_md gr.update。"""
            if not state or not state.get("scenario"):
                return gr.update(visible=False)
            return gr.update(
                visible=True,
                value=_build_todo_list(
                    state["scenario"],
                    state.get("done_steps", []),
                    state.get("current_step"),
                )
            )

        # ── 文字送出（streaming generator）──────────────────────────────────
        def on_submit(user_text, chat_history, log_history, state):
            if not user_text or not user_text.strip() or state is None:
                yield chat_history, log_history, state, "", gr.update(), gr.update()
                return

            chat_history = chat_history + [_user_msg(user_text)]
            cur = state["state"]

            if cur == "SCENARIO_SELECT":
                log_history = _log_append(log_history, "使用者選擇研究情境中……")
                yield chat_history, log_history, state, "", gr.update(), gr.update()
                reply, state = handle_scenario_select(user_text, state)
                if state["state"] == "UPLOAD_GUIDE":
                    log_history = _log_append(log_history, "✅ STEP 1 情境選擇完成")
                    chat_history = chat_history + [_assistant_msg(reply)]
                    yield chat_history, log_history, state, "", gr.update(), _todo_update(state)
                    # STEP 2
                    state["current_step"] = 2
                    log_history = _log_append(log_history, "⏳ STEP 2 上傳引導中……")
                    yield chat_history, log_history, state, "", gr.update(), _todo_update(state)
                    guide = get_upload_guide_message(state["scenario"])
                    chat_history = chat_history + [_assistant_msg(guide)]
                    state["done_steps"] = [1, 2]
                    state["current_step"] = 3
                    state["state"] = "UPLOAD_WAIT"
                    log_history = _log_append(log_history, "✅ STEP 2 完成，等待上傳文本")
                else:
                    chat_history = chat_history + [_assistant_msg(reply)]
                yield chat_history, log_history, state, "", gr.update(visible=(state["state"] == "UPLOAD_WAIT")), _todo_update(state)
                return

            if cur == "UPLOAD_WAIT":
                chat_history = chat_history + [_assistant_msg("請點擊下方「上傳已斷詞文本」選擇檔案並點擊「確認上傳」。")]
                yield chat_history, log_history, state, "", gr.update(visible=True), gr.update()
                return

            if cur == "COMPLETE":
                chat_history = chat_history + [_assistant_msg("分析已完成，請切換上方分頁查看各項結果。\n如需重新分析，請重新整理頁面。")]
                yield chat_history, log_history, state, "", gr.update(), gr.update()
                return

            chat_history = chat_history + [_assistant_msg(f"目前狀態：{STATES.get(cur, '未知')}，請依系統提示操作。")]
            yield chat_history, log_history, state, "", gr.update(), gr.update()

        send_btn.click(
            fn=on_submit,
            inputs=[msg_input, chatbot, exec_log, session_state],
            outputs=[chatbot, exec_log, session_state, msg_input, upload_row, todo_md],
        ).then(
            fn=_populate_tabs,
            inputs=[session_state],
            outputs=_TAB_OUTPUTS,
        )
        msg_input.submit(
            fn=on_submit,
            inputs=[msg_input, chatbot, exec_log, session_state],
            outputs=[chatbot, exec_log, session_state, msg_input, upload_row, todo_md],
        ).then(
            fn=_populate_tabs,
            inputs=[session_state],
            outputs=_TAB_OUTPUTS,
        )

        # ── 檔案上傳確認（streaming generator）──────────────────────────────
        def on_upload_confirm(files, chat_history, log_history, state):
            if not files or state is None:
                yield chat_history, log_history, state, gr.update(visible=True), gr.update()
                return

            # STEP 3（開始處理即關閉上傳彈窗，避免遮罩擋住操作）
            state["current_step"] = 3
            log_history = _log_append(log_history, "⏳ STEP 3 接收上傳檔案中……")
            yield chat_history, log_history, state, gr.update(visible=False), _todo_update(state)

            recv_msg, state = handle_uploaded_files(files, state)
            chat_history = chat_history + [_assistant_msg(recv_msg)]
            state["done_steps"] = list(set(state.get("done_steps", []) + [3]))
            state["current_step"] = 4
            log_history = _log_append(log_history, f"✅ STEP 3 已接收 {len(files)} 個檔案")
            yield chat_history, log_history, state, gr.update(visible=False), _todo_update(state)

            # STEP 4
            log_history = _log_append(log_history, "⏳ STEP 4 格式驗證中……")
            yield chat_history, log_history, state, gr.update(visible=False), _todo_update(state)

            check_msg, state, has_errors = run_format_check(state)
            chat_history = chat_history + [_assistant_msg(check_msg)]

            if has_errors:
                log_history = _log_append(log_history, "⚠️ STEP 4 格式不符，請重新上傳")
                state["current_step"] = 3
                yield chat_history, log_history, state, gr.update(visible=False), _todo_update(state)
                return

            state["done_steps"] = list(set(state.get("done_steps", []) + [4]))
            state["current_step"] = 5
            log_history = _log_append(log_history, "✅ STEP 4 格式驗證通過")
            # 格式通過後：顯示 metadata 說明表單，等使用者填寫後再開始分析
            state["state"] = "METADATA_INPUT"
            meta_hint = "請填寫語料 Metadata 說明（原文資訊、群組層與個體層），完成後點「確認並開始分析」。"
            chat_history = chat_history + [_assistant_msg(meta_hint)]
            log_history = _log_append(log_history, "📑 請填寫語料 Metadata 說明")
            yield chat_history, log_history, state, gr.update(visible=False), _todo_update(state)

        # ── metadata 表單 helpers ───────────────────────────────────────────
        def _build_meta_tables(state):
            import os
            files = [os.path.basename(p) for p in state.get("cleaned_files", [])]
            groups = sorted(set(f[0].upper() for f in files if f))
            return [[g, "", "", "", ""] for g in groups], [[f, "", "", "", "", ""] for f in files]

        def _df_rows(dfval):
            try:
                if isinstance(dfval, pd.DataFrame):
                    return dfval.fillna("").values.tolist()
                if isinstance(dfval, dict) and "data" in dfval:
                    return dfval["data"]
                if isinstance(dfval, list):
                    return dfval
            except Exception:
                pass
            return []

        def _compile_metadata(state, title, author, lang, pubyear, version,
                              ngroups, method, group_df, indiv_df):
            sc = state.get("scenario", "")
            L = ["## 語料 Metadata", "",
                 "### 一、原文資訊",
                 f"- 原文標題：{title or '（未填）'}",
                 f"- 作者：{author or '（未填）'}",
                 f"- 語言：{lang or '（未填）'}",
                 f"- 出版年：{pubyear or '（未填）'}",
                 f"- 版本：{version or '（未填）'}", ""]
            if sc in ("3", "4"):
                L += ["### 二、比較結構",
                      f"- 群組數量：{ngroups}",
                      f"- 比較方式：{method}", ""]
            gh = ["群組標籤", "翻譯類型", "翻譯方向", "譯者背景", "目標讀者"]
            L += ["### 三、群組層資訊"]
            for r in _df_rows(group_df):
                if not any(str(c).strip() for c in r):
                    continue
                L.append("- " + "；".join(f"{h}：{str(v).strip()}" for h, v in zip(gh, r) if str(v).strip()))
            ih = ["檔案", "譯者／系統名稱", "出版年", "譯者母語", "後編輯程度", "重譯序位"]
            L += ["", "### 四、個體層資訊"]
            for r in _df_rows(indiv_df):
                if not any(str(c).strip() for c in r):
                    continue
                L.append("- " + "；".join(f"{h}：{str(v).strip()}" for h, v in zip(ih, r) if str(v).strip()))
            return "\n".join(L)

        def _run_computation(chat_history, log_history, state):
            log_history = _log_append(log_history, "⏳ STEP 5 計算語言指標中……")
            yield chat_history, log_history, state, _todo_update(state)
            for prog_msg, is_final, state in run_computation_pipeline_streaming(state):
                if "指標計算完成" in prog_msg:
                    state["done_steps"] = list(set(state.get("done_steps", []) + [5])); state["current_step"] = 6
                elif "步驟 2/3" in prog_msg or "統計分析" in prog_msg and "完成" not in prog_msg:
                    state["current_step"] = 6
                elif "統計分析完成" in prog_msg or ("統計" in prog_msg and "完成" in prog_msg):
                    state["done_steps"] = list(set(state.get("done_steps", []) + [6])); state["current_step"] = 7
                elif "步驟 3/3" in prog_msg or "報告" in prog_msg and "完成" not in prog_msg:
                    state["current_step"] = 7
                elif is_final:
                    state["done_steps"] = list(set(state.get("done_steps", []) + [7, 8])); state["current_step"] = None
                log_history = _log_append(log_history, prog_msg)
                if is_final:
                    chat_history = chat_history + [_assistant_msg(prog_msg)]
                yield chat_history, log_history, state, _todo_update(state)

        def _show_metadata_form(state):
            _NO = gr.update()
            if not state or state.get("state") != "METADATA_INPUT":
                return _NO, _NO, _NO, _NO
            group_rows, indiv_rows = _build_meta_tables(state)
            show_cmp = state.get("scenario", "") in ("3", "4")
            return (gr.update(visible=True), gr.update(value=group_rows),
                    gr.update(value=indiv_rows), gr.update(visible=show_cmp))

        upload_btn.click(
            fn=on_upload_confirm,
            inputs=[file_input, chatbot, exec_log, session_state],
            outputs=[chatbot, exec_log, session_state, upload_row, todo_md],
        ).then(
            fn=_show_metadata_form,
            inputs=[session_state],
            outputs=[metadata_row, meta_group_df, meta_indiv_df, meta_compare_group],
        )

        cancel_btn.click(
            fn=lambda: (gr.update(visible=False), None),
            inputs=None,
            outputs=[upload_row, file_input],
        )

        def on_metadata_confirm(title, author, lang, pubyear, version,
                                ngroups, method, group_df, indiv_df,
                                chat_history, log_history, state):
            if state is None:
                yield chat_history, log_history, state, gr.update(), _todo_update(state)
                return
            state["user_metadata"] = _compile_metadata(state, title, author, lang, pubyear,
                                                        version, ngroups, method, group_df, indiv_df)
            log_history = _log_append(log_history, "✅ 已記錄語料 Metadata，開始分析")
            yield chat_history, log_history, state, gr.update(visible=False), _todo_update(state)
            for ch, lh, st, todo in _run_computation(chat_history, log_history, state):
                chat_history, log_history, state = ch, lh, st
                yield chat_history, log_history, state, gr.update(visible=False), todo

        meta_confirm_btn.click(
            fn=on_metadata_confirm,
            inputs=[meta_title, meta_author, meta_lang, meta_pubyear, meta_version,
                    meta_ngroups, meta_method, meta_group_df, meta_indiv_df,
                    chatbot, exec_log, session_state],
            outputs=[chatbot, exec_log, session_state, metadata_row, todo_md],
        ).then(fn=_populate_tabs, inputs=[session_state], outputs=_TAB_OUTPUTS)

        def on_metadata_skip(chat_history, log_history, state):
            if state is None:
                yield chat_history, log_history, state, gr.update(), _todo_update(state)
                return
            state["user_metadata"] = state.get("user_metadata", "") or "（使用者未提供語料 Metadata 說明）"
            log_history = _log_append(log_history, "⏭️ 已略過 Metadata，直接開始分析")
            yield chat_history, log_history, state, gr.update(visible=False), _todo_update(state)
            for ch, lh, st, todo in _run_computation(chat_history, log_history, state):
                chat_history, log_history, state = ch, lh, st
                yield chat_history, log_history, state, gr.update(visible=False), todo

        meta_cancel_btn.click(
            fn=on_metadata_skip,
            inputs=[chatbot, exec_log, session_state],
            outputs=[chatbot, exec_log, session_state, metadata_row, todo_md],
        ).then(fn=_populate_tabs, inputs=[session_state], outputs=_TAB_OUTPUTS)

    return demo


# ============================================================================
# SECTION 11: Entry Point
# ============================================================================

if __name__ == "__main__":
    print("=" * 60)
    print("重譯文本分析系統 AI Agent 啟動中……")
    print(f"主 Agent  : {LLM1_MODEL}")
    print(f"工作程序  : {LLM2_MODEL}")
    print(f"BASE_DIR  : {BASE_DIR}")
    print("=" * 60)
    app = create_app()
    app.launch(
        server_name="0.0.0.0",
        server_port=7860,
        share=False,
        inbrowser=False,
        allowed_paths=[str(SESSIONS_DIR)],
    )
